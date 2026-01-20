from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File, Depends, Response, Request
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
import bcrypt
import jwt
from io import BytesIO
import aiofiles
import base64
import httpx
import asyncio
import random
import string
import re

# Scheduler imports
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from apscheduler.triggers.cron import CronTrigger

# Job Scraper
from utils.job_scraper import job_scraper

# Email (Resend)
import resend

# Document processing
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from PyPDF2 import PdfReader

# AI Integration
from emergentintegrations.llm.chat import LlmChat, UserMessage

ROOT_DIR = Path(__file__).parent
# Load .env file only if it exists (for local development)
# In production, environment variables are injected by Kubernetes
env_file = ROOT_DIR / '.env'
if env_file.exists():
    load_dotenv(env_file, override=False)

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET_KEY', 'default-secret-key')
JWT_ALGORITHM = os.environ.get('JWT_ALGORITHM', 'HS256')
JWT_EXPIRATION = int(os.environ.get('JWT_EXPIRATION_HOURS', '24'))

# Emergent LLM Key
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY')

# RapidAPI Configuration for JSearch
RAPIDAPI_KEY = os.environ.get('RAPIDAPI_KEY')
RAPIDAPI_HOST = os.environ.get('RAPIDAPI_HOST', 'jsearch.p.rapidapi.com')

# LinkedIn OAuth Configuration
LINKEDIN_CLIENT_ID = os.environ.get('LINKEDIN_CLIENT_ID')
LINKEDIN_CLIENT_SECRET = os.environ.get('LINKEDIN_CLIENT_SECRET')

# Resend Email Configuration
RESEND_API_KEY = os.environ.get('RESEND_API_KEY')
SENDER_EMAIL = os.environ.get('SENDER_EMAIL', 'onboarding@resend.dev')
if RESEND_API_KEY:
    resend.api_key = RESEND_API_KEY

app = FastAPI(title="AI Resume Tailor API")
api_router = APIRouter(prefix="/api")

# Initialize the scheduler
scheduler = AsyncIOScheduler()

# ============ MODELS ============

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    name: str
    primary_technology: Optional[str] = ""
    sub_technologies: List[str] = []
    phone: Optional[str] = None
    location: Optional[str] = None

class UserLogin(BaseModel):
    email: EmailStr
    password: str

# OTP Verification Models
class SendOTPRequest(BaseModel):
    email: EmailStr
    name: str

class VerifyOTPRequest(BaseModel):
    email: EmailStr
    otp: str

class RegisterWithOTPRequest(BaseModel):
    email: EmailStr
    password: str
    name: str
    otp: str
    primary_technology: Optional[str] = ""
    sub_technologies: List[str] = []
    phone: Optional[str] = None
    location: Optional[str] = None

class UserProfileUpdate(BaseModel):
    name: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    primary_technology: Optional[str] = None
    sub_technologies: Optional[List[str]] = None
    profile_picture: Optional[str] = None
    linkedin_profile: Optional[str] = None
    salary_min: Optional[int] = None
    salary_max: Optional[int] = None
    salary_type: Optional[str] = None  # hourly, annual
    tax_types: Optional[List[str]] = None  # Fulltime, C2C, W2 Contract
    relocation_preference: Optional[str] = None  # yes, no, maybe
    location_preferences: Optional[List[str]] = None
    job_type_preferences: Optional[List[str]] = None  # remote, hybrid, onsite

class UserResponse(BaseModel):
    user_id: str
    email: str
    name: str
    primary_technology: str
    sub_technologies: List[str]
    phone: Optional[str]
    location: Optional[str]
    role: str
    created_at: datetime
    profile_picture: Optional[str] = None
    linkedin_profile: Optional[str] = None
    salary_min: Optional[int] = None
    salary_max: Optional[int] = None
    salary_type: Optional[str] = None
    tax_types: Optional[List[str]] = None
    relocation_preference: Optional[str] = None
    location_preferences: Optional[List[str]] = None
    job_type_preferences: Optional[List[str]] = None

class TokenResponse(BaseModel):
    access_token: str
    token_type: str
    user: UserResponse

class LinkedInCallbackRequest(BaseModel):
    code: str
    redirect_uri: str

class ResumeCreate(BaseModel):
    original_content: str
    file_name: str

class ResumeUpdate(BaseModel):
    tailored_content: Optional[str] = None
    target_job_title: Optional[str] = None
    target_job_description: Optional[str] = None

class JobPortalCreate(BaseModel):
    name: str
    url: str
    technology: str
    description: Optional[str] = None

class JobApplicationCreate(BaseModel):
    job_portal_id: str
    job_title: str
    job_description: str
    company_name: str
    resume_id: str
    cover_letter: Optional[str] = None
    tailored_content: Optional[str] = None  # The tailored resume content used for application
    job_source: Optional[str] = None  # Indeed, Dice, etc.
    apply_link: Optional[str] = None  # Job application URL

class EmailCreate(BaseModel):
    application_id: str
    subject: str
    content: str
    email_type: str  # sent, received, scheduled

class TailorResumeRequest(BaseModel):
    resume_id: str
    job_title: str
    job_description: str
    technologies: List[str] = []
    company_name: Optional[str] = ""
    custom_prompt: Optional[str] = None  # Custom AI command/prompt
    generate_versions: bool = False  # Optional: generate 2-3 ATS versions
    ats_optimize: bool = True  # ATS-friendly optimization

class OptimizeResumeRequest(BaseModel):
    target_role: str = ""  # Optional target role for optimization
    generate_versions: bool = False  # Generate 2-3 versions

class AnalyzeResumeRequest(BaseModel):
    resume_id: str

class CreateMasterResumeRequest(BaseModel):
    resume_id: str

class GenerateVersionsRequest(BaseModel):
    resume_id: str
    primary_technology: Optional[str] = None

class GenerateCoverLetterRequest(BaseModel):
    resume_id: str
    job_title: str
    company_name: str
    job_description: str

class EmailReplyRequest(BaseModel):
    original_email: str
    context: str
    tone: str = "professional"

# Auto-Apply Models
class AutoApplySettings(BaseModel):
    enabled: bool = False
    resume_id: str = ""
    job_keywords: List[str] = []
    locations: List[str] = ["United States"]
    employment_types: List[str] = ["FULL_TIME"]
    min_salary: Optional[int] = None
    max_applications_per_day: int = 10
    auto_tailor_resume: bool = True

class AutoApplySettingsUpdate(BaseModel):
    enabled: Optional[bool] = None
    resume_id: Optional[str] = None
    job_keywords: Optional[List[str]] = None
    locations: Optional[List[str]] = None
    employment_types: Optional[List[str]] = None
    min_salary: Optional[int] = None
    max_applications_per_day: Optional[int] = None
    auto_tailor_resume: Optional[bool] = None

# ============ EMAIL CENTER MODELS ============

class EmailAccountConnect(BaseModel):
    provider: str  # gmail, outlook, imap
    # For IMAP/SMTP
    imap_host: Optional[str] = None
    imap_port: Optional[int] = 993
    smtp_host: Optional[str] = None
    smtp_port: Optional[int] = 587
    email_address: Optional[str] = None
    password: Optional[str] = None  # App password for IMAP
    use_ssl: bool = True

class EmailAccountResponse(BaseModel):
    account_id: str
    provider: str
    email_address: str
    is_connected: bool
    is_primary: bool
    connected_at: datetime
    last_sync: Optional[datetime] = None

class InboxMessage(BaseModel):
    message_id: str
    from_email: str
    from_name: Optional[str] = None
    to_email: str
    subject: str
    body_preview: str
    body: Optional[str] = None
    received_at: datetime
    is_read: bool
    is_recruiter: bool = False
    labels: List[str] = []

class SendEmailRequest(BaseModel):
    to_addresses: List[EmailStr]
    subject: str
    body: str
    body_type: str = "text"  # text or html
    account_id: Optional[str] = None  # Use primary if not specified
    
class AIComposeRequest(BaseModel):
    job_title: str
    company_name: str
    job_description: str
    resume_id: str
    recipient_email: EmailStr
    tone: str = "professional"  # professional, friendly, formal
    
class AIReplyRequest(BaseModel):
    original_email: str
    original_subject: str
    sender_email: str
    context: str = ""
    tone: str = "professional"
    account_id: Optional[str] = None

class EmailCenterSettings(BaseModel):
    auto_reply_enabled: bool = False
    auto_apply_compose: bool = True
    reply_approval_required: bool = True  # Require approval before sending AI replies
    signature: Optional[str] = None

# ============ AUTH HELPERS ============

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_access_token(data: dict, expires_delta: timedelta = None) -> str:
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.now(timezone.utc) + expires_delta
    else:
        expire = datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION)
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(request: Request) -> dict:
    # Check cookie first
    token = request.cookies.get("session_token")
    
    # Then check Authorization header
    if not token:
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            token = auth_header[7:]
    
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    # Check if it's a session token (from Google OAuth)
    session = await db.user_sessions.find_one({"session_token": token}, {"_id": 0})
    if session:
        expires_at = session.get("expires_at")
        if isinstance(expires_at, str):
            expires_at = datetime.fromisoformat(expires_at)
        if expires_at.tzinfo is None:
            expires_at = expires_at.replace(tzinfo=timezone.utc)
        if expires_at < datetime.now(timezone.utc):
            raise HTTPException(status_code=401, detail="Session expired")
        
        user = await db.users.find_one({"user_id": session["user_id"]}, {"_id": 0})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        return user
    
    # Otherwise, decode JWT
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token")
        
        user = await db.users.find_one({"user_id": user_id}, {"_id": 0})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

async def get_admin_user(request: Request) -> dict:
    user = await get_current_user(request)
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return user

# ============ AUTH ROUTES ============

# OTP Helper Functions
def generate_otp() -> str:
    """Generate a 6-digit OTP"""
    return ''.join(random.choices(string.digits, k=6))


# API Health Check
@api_router.get("/health")
async def api_health_check():
    """API health check endpoint"""
    try:
        # Test database connection
        await db.command('ping')
        db_status = "connected"
    except Exception as e:
        db_status = f"error: {str(e)}"
    
    return {
        "status": "healthy",
        "service": "careerquest-api",
        "database": db_status
    }


@api_router.post("/auth/send-otp")
async def send_otp(data: SendOTPRequest):
    """Send OTP to email for verification"""
    # Check if email is already registered
    existing = await db.users.find_one({"email": data.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered. Please login instead.")
    
    # Generate OTP
    otp = generate_otp()
    expires_at = datetime.now(timezone.utc) + timedelta(minutes=10)
    
    # Store OTP in database (upsert to handle resends)
    await db.otp_verifications.update_one(
        {"email": data.email},
        {
            "$set": {
                "otp": otp,
                "name": data.name,
                "expires_at": expires_at.isoformat(),
                "verified": False,
                "created_at": datetime.now(timezone.utc).isoformat()
            }
        },
        upsert=True
    )
    
    # Return OTP directly (built-in verification system)
    logger.info(f"OTP generated for {data.email}: {otp}")
    
    return {
        "message": "Verification code generated",
        "email": data.email,
        "expires_in_minutes": 10,
        "otp": otp  # Return OTP for built-in verification
    }


@api_router.post("/auth/verify-otp")
async def verify_otp(data: VerifyOTPRequest):
    """Verify the OTP code"""
    # Find OTP record
    otp_record = await db.otp_verifications.find_one(
        {"email": data.email},
        {"_id": 0}
    )
    
    if not otp_record:
        raise HTTPException(status_code=400, detail="No verification code found. Please request a new code.")
    
    # Check expiration
    expires_at = datetime.fromisoformat(otp_record["expires_at"])
    if datetime.now(timezone.utc) > expires_at:
        raise HTTPException(status_code=400, detail="Verification code has expired. Please request a new code.")
    
    # Verify OTP
    if otp_record["otp"] != data.otp:
        raise HTTPException(status_code=400, detail="Invalid verification code. Please try again.")
    
    # Mark as verified
    await db.otp_verifications.update_one(
        {"email": data.email},
        {"$set": {"verified": True}}
    )
    
    return {
        "message": "Email verified successfully",
        "email": data.email,
        "verified": True
    }


@api_router.post("/auth/register-with-otp", response_model=TokenResponse)
async def register_with_otp(user_data: RegisterWithOTPRequest):
    """Register a new user after OTP verification"""
    # Check if email already registered
    existing = await db.users.find_one({"email": user_data.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Verify OTP was completed
    otp_record = await db.otp_verifications.find_one(
        {"email": user_data.email},
        {"_id": 0}
    )
    
    if not otp_record:
        raise HTTPException(status_code=400, detail="Please verify your email first")
    
    if not otp_record.get("verified"):
        # Verify the OTP inline
        if otp_record["otp"] != user_data.otp:
            raise HTTPException(status_code=400, detail="Invalid verification code")
        
        # Check expiration
        expires_at = datetime.fromisoformat(otp_record["expires_at"])
        if datetime.now(timezone.utc) > expires_at:
            raise HTTPException(status_code=400, detail="Verification code has expired")
    
    # Create user
    user_id = f"user_{uuid.uuid4().hex[:12]}"
    hashed_password = hash_password(user_data.password)
    
    user_doc = {
        "user_id": user_id,
        "email": user_data.email,
        "password": hashed_password,
        "name": user_data.name,
        "primary_technology": user_data.primary_technology,
        "sub_technologies": user_data.sub_technologies,
        "phone": user_data.phone,
        "location": user_data.location,
        "role": "candidate",
        "email_verified": True,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_doc)
    
    # Clean up OTP record
    await db.otp_verifications.delete_one({"email": user_data.email})
    
    token = create_access_token({"user_id": user_id, "email": user_data.email})
    
    user_response = UserResponse(
        user_id=user_id,
        email=user_data.email,
        name=user_data.name,
        primary_technology=user_data.primary_technology,
        sub_technologies=user_data.sub_technologies,
        phone=user_data.phone,
        location=user_data.location,
        role="candidate",
        created_at=datetime.now(timezone.utc)
    )
    
    return TokenResponse(access_token=token, token_type="bearer", user=user_response)


@api_router.post("/auth/resend-otp")
async def resend_otp(data: SendOTPRequest):
    """Resend OTP to email"""
    # Generate new OTP
    otp = generate_otp()
    expires_at = datetime.now(timezone.utc) + timedelta(minutes=10)
    
    # Update OTP in database
    result = await db.otp_verifications.update_one(
        {"email": data.email},
        {
            "$set": {
                "otp": otp,
                "expires_at": expires_at.isoformat(),
                "verified": False
            }
        }
    )
    
    if result.matched_count == 0:
        # Create new record if doesn't exist
        await db.otp_verifications.insert_one({
            "email": data.email,
            "name": data.name,
            "otp": otp,
            "expires_at": expires_at.isoformat(),
            "verified": False,
            "created_at": datetime.now(timezone.utc).isoformat()
        })
    
    # Return OTP directly (built-in verification system)
    logger.info(f"OTP regenerated for {data.email}: {otp}")
    
    return {
        "message": "New verification code generated",
        "email": data.email,
        "expires_in_minutes": 10,
        "otp": otp  # Return OTP for built-in verification
    }


# Keep old register endpoint for backward compatibility (can be removed later)
@api_router.post("/auth/register", response_model=TokenResponse)
async def register(user_data: UserCreate):
    existing = await db.users.find_one({"email": user_data.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user_id = f"user_{uuid.uuid4().hex[:12]}"
    hashed_password = hash_password(user_data.password)
    
    user_doc = {
        "user_id": user_id,
        "email": user_data.email,
        "password": hashed_password,
        "name": user_data.name,
        "primary_technology": user_data.primary_technology,
        "sub_technologies": user_data.sub_technologies,
        "phone": user_data.phone,
        "location": user_data.location,
        "role": "candidate",
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_doc)
    
    token = create_access_token({"user_id": user_id, "email": user_data.email})
    
    user_response = UserResponse(
        user_id=user_id,
        email=user_data.email,
        name=user_data.name,
        primary_technology=user_data.primary_technology,
        sub_technologies=user_data.sub_technologies,
        phone=user_data.phone,
        location=user_data.location,
        role="candidate",
        created_at=datetime.now(timezone.utc)
    )
    
    return TokenResponse(access_token=token, token_type="bearer", user=user_response)

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(credentials: UserLogin, response: Response):
    user = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if not user or not verify_password(credentials.password, user["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    token = create_access_token({"user_id": user["user_id"], "email": user["email"]})
    
    # Set cookie
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRATION * 3600
    )
    
    created_at = user.get("created_at")
    if isinstance(created_at, str):
        created_at = datetime.fromisoformat(created_at)
    
    user_response = UserResponse(
        user_id=user["user_id"],
        email=user["email"],
        name=user["name"],
        primary_technology=user.get("primary_technology", ""),
        sub_technologies=user.get("sub_technologies", []),
        phone=user.get("phone"),
        location=user.get("location"),
        role=user.get("role", "candidate"),
        created_at=created_at
    )
    
    return TokenResponse(access_token=token, token_type="bearer", user=user_response)

@api_router.get("/auth/me")
async def get_me(request: Request):
    user = await get_current_user(request)
    created_at = user.get("created_at")
    if isinstance(created_at, str):
        created_at = datetime.fromisoformat(created_at)
    
    return {
        "user_id": user["user_id"],
        "email": user["email"],
        "name": user["name"],
        "primary_technology": user.get("primary_technology", ""),
        "sub_technologies": user.get("sub_technologies", []),
        "phone": user.get("phone"),
        "location": user.get("location"),
        "role": user.get("role", "candidate"),
        "picture": user.get("picture"),
        "profile_picture": user.get("profile_picture"),
        "linkedin_profile": user.get("linkedin_profile"),
        "salary_min": user.get("salary_min"),
        "salary_max": user.get("salary_max"),
        "salary_type": user.get("salary_type"),
        "tax_types": user.get("tax_types", []),
        "relocation_preference": user.get("relocation_preference"),
        "location_preferences": user.get("location_preferences", []),
        "job_type_preferences": user.get("job_type_preferences", []),
        "created_at": created_at.isoformat() if created_at else None
    }

@api_router.put("/auth/profile")
async def update_profile(data: UserProfileUpdate, request: Request):
    user = await get_current_user(request)
    
    update_data = {}
    if data.name is not None:
        update_data["name"] = data.name
    if data.phone is not None:
        update_data["phone"] = data.phone
    if data.location is not None:
        update_data["location"] = data.location
    if data.primary_technology is not None:
        update_data["primary_technology"] = data.primary_technology
    if data.sub_technologies is not None:
        update_data["sub_technologies"] = data.sub_technologies
    if data.profile_picture is not None:
        update_data["profile_picture"] = data.profile_picture
    if data.linkedin_profile is not None:
        update_data["linkedin_profile"] = data.linkedin_profile
    if data.salary_min is not None:
        update_data["salary_min"] = data.salary_min
    if data.salary_max is not None:
        update_data["salary_max"] = data.salary_max
    if data.salary_type is not None:
        update_data["salary_type"] = data.salary_type
    if data.tax_types is not None:
        update_data["tax_types"] = data.tax_types
    if data.relocation_preference is not None:
        update_data["relocation_preference"] = data.relocation_preference
    if data.location_preferences is not None:
        update_data["location_preferences"] = data.location_preferences
    if data.job_type_preferences is not None:
        update_data["job_type_preferences"] = data.job_type_preferences
    
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.users.update_one(
        {"user_id": user["user_id"]},
        {"$set": update_data}
    )
    
    # Fetch updated user
    updated_user = await db.users.find_one({"user_id": user["user_id"]}, {"_id": 0, "password_hash": 0, "password": 0})
    
    return {"message": "Profile updated successfully", "user": updated_user}


@api_router.post("/auth/profile-photo")
async def upload_profile_photo(request: Request, file: UploadFile = File(...)):
    """Upload and update user profile photo."""
    user = await get_current_user(request)
    
    # Validate file type
    allowed_types = ['image/jpeg', 'image/png', 'image/gif', 'image/webp']
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=400, 
            detail="Invalid file type. Please upload a JPEG, PNG, GIF, or WebP image."
        )
    
    # Validate file size (max 5MB)
    contents = await file.read()
    if len(contents) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File size must be less than 5MB")
    
    # Convert to base64 data URL for storage
    import base64
    file_extension = file.content_type.split('/')[-1]
    base64_data = base64.b64encode(contents).decode('utf-8')
    data_url = f"data:{file.content_type};base64,{base64_data}"
    
    # Update user profile with the photo
    await db.users.update_one(
        {"user_id": user["user_id"]},
        {
            "$set": {
                "profile_picture": data_url,
                "picture": data_url,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }
        }
    )
    
    # Fetch updated user
    updated_user = await db.users.find_one({"user_id": user["user_id"]}, {"_id": 0, "password": 0})
    
    return {
        "message": "Profile photo updated successfully",
        "profile_picture": data_url,
        "user": updated_user
    }


@api_router.delete("/auth/profile-photo")
async def delete_profile_photo(request: Request):
    """Remove user profile photo."""
    user = await get_current_user(request)
    
    await db.users.update_one(
        {"user_id": user["user_id"]},
        {
            "$set": {
                "profile_picture": None,
                "picture": None,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }
        }
    )
    
    return {"message": "Profile photo removed successfully"}


@api_router.get("/auth/profile-completeness")
async def get_profile_completeness(request: Request):
    user = await get_current_user(request)
    
    # Define required fields and their weights
    fields = {
        "name": 10,
        "email": 10,
        "primary_technology": 15,
        "sub_technologies": 10,
        "location": 10,
        "phone": 5,
        "linkedin_profile": 10,
        "salary_min": 5,
        "salary_max": 5,
        "tax_types": 5,
        "relocation_preference": 5,
        "location_preferences": 5,
        "job_type_preferences": 5,
    }
    
    total_weight = sum(fields.values())
    completed_weight = 0
    missing_fields = []
    
    for field, weight in fields.items():
        value = user.get(field)
        if value and (not isinstance(value, list) or len(value) > 0):
            completed_weight += weight
        else:
            missing_fields.append(field)
    
    # Check if user has uploaded a resume
    resume_count = await db.resumes.count_documents({"user_id": user["user_id"]})
    if resume_count > 0:
        completed_weight += 10  # Resume adds 10%
    else:
        missing_fields.append("resume")
    
    total_weight += 10  # Add resume weight to total
    
    percentage = int((completed_weight / total_weight) * 100)
    
    return {
        "percentage": percentage,
        "completed_weight": completed_weight,
        "total_weight": total_weight,
        "missing_fields": missing_fields
    }

@api_router.post("/auth/logout")
async def logout(request: Request, response: Response):
    token = request.cookies.get("session_token")
    if token:
        await db.user_sessions.delete_one({"session_token": token})
    response.delete_cookie(key="session_token", path="/")
    return {"message": "Logged out successfully"}

# Google OAuth Session endpoint
@api_router.post("/auth/session")
async def create_session(request: Request, response: Response):
    session_id = request.headers.get("X-Session-ID")
    logger.info(f"Session endpoint called with session_id: {session_id[:20] if session_id else 'None'}...")
    
    if not session_id:
        logger.error("No session ID provided in request")
        raise HTTPException(status_code=400, detail="Session ID required")
    
    # Fetch user data from Emergent Auth
    import httpx
    try:
        async with httpx.AsyncClient(timeout=30.0) as http_client:
            logger.info("Fetching session data from Emergent Auth...")
            auth_response = await http_client.get(
                "https://demobackend.emergentagent.com/auth/v1/env/oauth/session-data",
                headers={"X-Session-ID": session_id}
            )
            logger.info(f"Emergent Auth response status: {auth_response.status_code}")
            
            if auth_response.status_code != 200:
                error_detail = auth_response.text
                logger.error(f"Emergent Auth error: {error_detail}")
                # Return more specific error to help debug
                raise HTTPException(
                    status_code=401, 
                    detail=f"Session expired or invalid. Please try signing in again."
                )
            
            auth_data = auth_response.json()
            logger.info(f"Auth data received for email: {auth_data.get('email')}")
    except httpx.TimeoutException:
        logger.error("Timeout while fetching session data from Emergent Auth")
        raise HTTPException(status_code=504, detail="Authentication service timeout. Please try again.")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching session data: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Authentication error. Please try again.")
    
    email = auth_data.get("email")
    name = auth_data.get("name")
    picture = auth_data.get("picture")
    session_token = auth_data.get("session_token")
    
    # Check if user exists
    user = await db.users.find_one({"email": email}, {"_id": 0})
    
    if not user:
        # Create new user
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        user = {
            "user_id": user_id,
            "email": email,
            "name": name,
            "picture": picture,
            "primary_technology": "",
            "sub_technologies": [],
            "role": "candidate",
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.users.insert_one(user)
    else:
        user_id = user["user_id"]
        # Update user info
        await db.users.update_one(
            {"user_id": user_id},
            {"$set": {"name": name, "picture": picture}}
        )
    
    # Store session
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": session_token,
        "expires_at": expires_at.isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    # Create JWT access token for API authorization
    access_token_expires = timedelta(days=7)
    access_token = create_access_token(
        data={"user_id": user_id, "email": email},
        expires_delta=access_token_expires
    )
    
    # Set cookie
    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=7 * 24 * 3600
    )
    
    # Return user data with access token
    return {
        "user_id": user_id,
        "email": email,
        "name": name,
        "picture": picture,
        "role": user.get("role", "candidate"),
        "access_token": access_token,
        "token_type": "bearer"
    }

# LinkedIn OAuth endpoint
@api_router.post("/auth/linkedin/callback", response_model=TokenResponse)
async def linkedin_callback(data: LinkedInCallbackRequest):
    """
    Handle LinkedIn OAuth callback - exchange code for access token and get user profile.
    """
    if not LINKEDIN_CLIENT_ID or not LINKEDIN_CLIENT_SECRET:
        raise HTTPException(
            status_code=500, 
            detail="LinkedIn OAuth is not configured. Please add LINKEDIN_CLIENT_ID and LINKEDIN_CLIENT_SECRET to your environment."
        )
    
    try:
        async with httpx.AsyncClient() as http_client:
            # Step 1: Exchange authorization code for access token
            token_response = await http_client.post(
                "https://www.linkedin.com/oauth/v2/accessToken",
                data={
                    "grant_type": "authorization_code",
                    "code": data.code,
                    "redirect_uri": data.redirect_uri,
                    "client_id": LINKEDIN_CLIENT_ID,
                    "client_secret": LINKEDIN_CLIENT_SECRET,
                },
                headers={"Content-Type": "application/x-www-form-urlencoded"},
                timeout=15.0
            )
            
            if token_response.status_code != 200:
                error_data = token_response.json()
                logger.error(f"LinkedIn token exchange failed: {error_data}")
                raise HTTPException(status_code=400, detail=f"LinkedIn authentication failed: {error_data.get('error_description', 'Unknown error')}")
            
            token_data = token_response.json()
            access_token = token_data.get("access_token")
            
            # Step 2: Get user profile using OpenID Connect userinfo endpoint
            profile_response = await http_client.get(
                "https://api.linkedin.com/v2/userinfo",
                headers={"Authorization": f"Bearer {access_token}"},
                timeout=10.0
            )
            
            if profile_response.status_code != 200:
                logger.error(f"LinkedIn profile fetch failed: {profile_response.text}")
                raise HTTPException(status_code=400, detail="Failed to fetch LinkedIn profile")
            
            profile_data = profile_response.json()
            
            # Extract user information
            linkedin_id = profile_data.get("sub")
            email = profile_data.get("email")
            name = profile_data.get("name")
            picture = profile_data.get("picture")
            
            if not email:
                raise HTTPException(status_code=400, detail="LinkedIn account does not have an email address")
            
            # Check if user exists by LinkedIn ID or email
            user = await db.users.find_one(
                {"$or": [{"linkedin_id": linkedin_id}, {"email": email}]},
                {"_id": 0}
            )
            
            if user:
                # Update existing user with LinkedIn info
                await db.users.update_one(
                    {"user_id": user["user_id"]},
                    {"$set": {
                        "linkedin_id": linkedin_id,
                        "picture": picture or user.get("picture"),
                        "profile_picture": picture or user.get("profile_picture"),
                        "last_login": datetime.now(timezone.utc).isoformat()
                    }}
                )
                user_id = user["user_id"]
            else:
                # Create new user
                user_id = f"user_{uuid.uuid4().hex[:12]}"
                user = {
                    "user_id": user_id,
                    "email": email,
                    "name": name,
                    "linkedin_id": linkedin_id,
                    "picture": picture,
                    "profile_picture": picture,
                    "primary_technology": "",
                    "sub_technologies": [],
                    "role": "candidate",
                    "created_at": datetime.now(timezone.utc).isoformat(),
                    "last_login": datetime.now(timezone.utc).isoformat()
                }
                await db.users.insert_one(user)
            
            # Fetch updated user
            user = await db.users.find_one({"user_id": user_id}, {"_id": 0, "password_hash": 0})
            
            # Generate JWT token
            token_payload = {
                "user_id": user_id,
                "email": email,
                "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION)
            }
            access_token_jwt = jwt.encode(token_payload, JWT_SECRET, algorithm=JWT_ALGORITHM)
            
            return TokenResponse(
                access_token=access_token_jwt,
                token_type="bearer",
                user=UserResponse(
                    user_id=user["user_id"],
                    email=user["email"],
                    name=user["name"],
                    primary_technology=user.get("primary_technology", ""),
                    sub_technologies=user.get("sub_technologies", []),
                    phone=user.get("phone"),
                    location=user.get("location"),
                    role=user.get("role", "candidate"),
                    created_at=datetime.fromisoformat(user["created_at"]) if isinstance(user.get("created_at"), str) else user.get("created_at"),
                    profile_picture=user.get("profile_picture") or user.get("picture"),
                    linkedin_profile=user.get("linkedin_profile"),
                    salary_min=user.get("salary_min"),
                    salary_max=user.get("salary_max"),
                    salary_type=user.get("salary_type"),
                    tax_type=user.get("tax_type"),
                    relocation_preference=user.get("relocation_preference"),
                    location_preferences=user.get("location_preferences", []),
                    job_type_preferences=user.get("job_type_preferences", [])
                )
            )
            
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="LinkedIn authentication timed out")
    except Exception as e:
        logger.error(f"LinkedIn authentication error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"LinkedIn authentication failed: {str(e)}")

# ============ RESUME ROUTES ============

async def extract_profile_from_resume(text_content: str, user_id: str):
    """Extract profile details from resume and update user profile"""
    
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"extract_profile_{uuid.uuid4().hex[:8]}",
            system_message="""You are an expert at extracting structured data from resumes. Extract contact and professional information accurately."""
        ).with_model("openai", "gpt-5.2")
        
        extraction_prompt = f"""Extract the following information from this resume. Return ONLY valid JSON.

RESUME:
{text_content}

Extract and return JSON format (use null for missing fields):
{{
    "name": "<full name>",
    "email": "<email address or null>",
    "phone": "<phone number or null>",
    "location": "<city, state or full address or null>",
    "linkedin_profile": "<linkedin URL or null>",
    "primary_technology": "<main technology/skill area: Java, Python, PHP, AI, React, Frontend, Backend, Full Stack, DevOps, Mobile, or null>",
    "sub_technologies": ["<skill1>", "<skill2>", "<skill3>", ...],
    "years_of_experience": <number or null>,
    "current_job_title": "<most recent job title or null>",
    "current_company": "<most recent company or null>",
    "education": "<highest degree and institution or null>",
    "certifications": ["<cert1>", "<cert2>", ...]
}}

Rules:
- For primary_technology, choose the BEST match from: Java, Python, PHP, AI, React, Frontend, Backend, Full Stack, DevOps, Mobile
- For sub_technologies, list specific tools, frameworks, and languages mentioned
- Extract phone in any format found
- Return null (not empty string) for missing fields"""

        extraction_message = UserMessage(text=extraction_prompt)
        extraction_response = await chat.send_message(extraction_message)
        
        # Parse the JSON response
        import json
        response_text = extraction_response.strip()
        if response_text.startswith('```'):
            response_text = response_text.split('\n', 1)[1]
            if '```' in response_text:
                response_text = response_text.split('```')[0]
        
        extracted_data = json.loads(response_text)
        logger.info(f"Extracted profile data: {extracted_data}")
        
        # Build update document (only update non-null fields)
        update_fields = {}
        
        if extracted_data.get("name"):
            update_fields["name"] = extracted_data["name"]
        if extracted_data.get("phone"):
            update_fields["phone"] = extracted_data["phone"]
        if extracted_data.get("location"):
            update_fields["location"] = extracted_data["location"]
        if extracted_data.get("linkedin_profile"):
            update_fields["linkedin_profile"] = extracted_data["linkedin_profile"]
        if extracted_data.get("primary_technology"):
            update_fields["primary_technology"] = extracted_data["primary_technology"]
        if extracted_data.get("sub_technologies") and len(extracted_data["sub_technologies"]) > 0:
            update_fields["sub_technologies"] = extracted_data["sub_technologies"]
        if extracted_data.get("years_of_experience"):
            update_fields["years_of_experience"] = extracted_data["years_of_experience"]
        if extracted_data.get("current_job_title"):
            update_fields["current_job_title"] = extracted_data["current_job_title"]
        if extracted_data.get("current_company"):
            update_fields["current_company"] = extracted_data["current_company"]
        if extracted_data.get("education"):
            update_fields["education"] = extracted_data["education"]
        if extracted_data.get("certifications") and len(extracted_data["certifications"]) > 0:
            update_fields["certifications"] = extracted_data["certifications"]
        
        if update_fields:
            update_fields["profile_extracted_from_resume"] = True
            update_fields["profile_extracted_at"] = datetime.now(timezone.utc).isoformat()
            update_fields["updated_at"] = datetime.now(timezone.utc).isoformat()
            
            await db.users.update_one(
                {"user_id": user_id},
                {"$set": update_fields}
            )
            logger.info(f"Profile updated for user {user_id} with {len(update_fields)} fields from resume")
        
        return extracted_data
        
    except Exception as e:
        logger.error(f"Error extracting profile from resume: {str(e)}")
        return None

async def auto_analyze_resume(resume_id: str, text_content: str, user_id: str):
    """Automatically analyze resume, extract profile, create master, and generate versions after upload"""
    
    # First, extract profile details and update user
    extracted_profile = await extract_profile_from_resume(text_content, user_id)
    
    # Get updated user profile for technology info
    user_profile = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    primary_tech = user_profile.get("primary_technology", "") if user_profile else ""
    sub_techs = user_profile.get("sub_technologies", []) if user_profile else []
    
    results = {
        "analysis": None,
        "master_resume": None,
        "title_versions": [],
        "extracted_profile": extracted_profile
    }
    
    try:
        # Step 1: Analyze Resume
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"auto_analyze_{resume_id}_{uuid.uuid4().hex[:8]}",
            system_message="""You are an expert resume analyst and career consultant. Analyze resumes thoroughly and provide actionable feedback."""
        ).with_model("openai", "gpt-5.2")
        
        analysis_prompt = f"""Analyze this resume and provide a comprehensive assessment.

RESUME:
{text_content}

Provide your analysis in the following JSON format (return ONLY valid JSON, no markdown):
{{
    "score": <number 0-100>,
    "grade": "<A/B/C/D/F>",
    "summary": "<2-3 sentence overall assessment>",
    "missing_info": {{
        "phone": <true if missing, false if present>,
        "email": <true if missing, false if present>,
        "address_location": <true if missing, false if present>,
        "linkedin": <true if missing, false if present>,
        "professional_summary": <true if missing, false if present>,
        "skills_section": <true if missing, false if present>,
        "education": <true if missing, false if present>,
        "work_experience": <true if missing, false if present>,
        "quantifiable_achievements": <true if missing/weak, false if strong>
    }},
    "strengths": ["<strength 1>", "<strength 2>", "<strength 3>"],
    "weaknesses": ["<weakness 1>", "<weakness 2>", "<weakness 3>"],
    "improvement_suggestions": ["<suggestion 1>", "<suggestion 2>", "<suggestion 3>", "<suggestion 4>"],
    "ats_compatibility": {{
        "score": <number 0-100>,
        "issues": ["<issue 1>", "<issue 2>"]
    }},
    "detected_skills": ["<skill1>", "<skill2>", ...],
    "detected_job_titles": ["<title1>", "<title2>", ...],
    "experience_level": "<Entry/Mid/Senior/Executive>"
}}"""

        analysis_message = UserMessage(text=analysis_prompt)
        analysis_response = await chat.send_message(analysis_message)
        
        # Parse the JSON response
        try:
            import json
            response_text = analysis_response.strip()
            if response_text.startswith('```'):
                response_text = response_text.split('\n', 1)[1]
                if response_text.endswith('```'):
                    response_text = response_text[:-3]
            results["analysis"] = json.loads(response_text)
        except:
            results["analysis"] = {
                "score": 50, "grade": "C",
                "summary": "Analysis completed with limited parsing.",
                "missing_info": {}, "strengths": [], "weaknesses": [],
                "improvement_suggestions": [], "ats_compatibility": {"score": 50, "issues": []},
                "detected_skills": [], "detected_job_titles": [], "experience_level": "Unknown"
            }
        
        logger.info(f"Resume {resume_id} analyzed: Score {results['analysis'].get('score', 'N/A')}")
        
        # Step 2: Create Master Resume
        master_chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"auto_master_{resume_id}_{uuid.uuid4().hex[:8]}",
            system_message="""You are an expert resume writer. Create polished, ATS-friendly resumes."""
        ).with_model("openai", "gpt-5.2")
        
        master_prompt = f"""Transform this resume into a polished MASTER RESUME.

ORIGINAL RESUME:
{text_content}

USER'S PRIMARY TECHNOLOGY: {primary_tech or 'Not specified'}
USER'S SKILLS: {', '.join(sub_techs) if sub_techs else 'Not specified'}

CREATE A MASTER RESUME that:
1. PROFESSIONAL SUMMARY: Write a compelling 3-4 line summary
2. FIX ALL ISSUES: Grammar, spelling, formatting
3. SKILLS SECTION: Organize into categories (Technical, Tools, Soft Skills)
4. EXPERIENCE: Rewrite with action verbs and quantifiable metrics
5. FORMATTING: Use standard ATS-friendly sections
6. CONTACT INFO: Ensure placeholders for phone, email, location, LinkedIn

Return ONLY the master resume content in plain text format."""

        master_message = UserMessage(text=master_prompt)
        results["master_resume"] = await master_chat.send_message(master_message)
        
        logger.info(f"Master resume created for {resume_id}")
        
        # Step 3: Generate Title Versions
        # Determine technology for title generation
        detected_tech = primary_tech
        if not detected_tech and results["analysis"].get("detected_skills"):
            skills = results["analysis"]["detected_skills"]
            if any(s.lower() in ['react', 'react.js', 'reactjs'] for s in skills):
                detected_tech = "React"
            elif any(s.lower() in ['python', 'django', 'flask'] for s in skills):
                detected_tech = "Python"
            elif any(s.lower() in ['java', 'spring', 'springboot'] for s in skills):
                detected_tech = "Java"
            elif any(s.lower() in ['node', 'node.js', 'nodejs', 'express'] for s in skills):
                detected_tech = "Backend"
            elif any(s.lower() in ['ai', 'ml', 'machine learning', 'tensorflow', 'pytorch'] for s in skills):
                detected_tech = "AI"
        
        tech_titles = {
            "Java": ["Java Developer", "Java Software Engineer", "Backend Developer", "Full Stack Java Developer"],
            "Python": ["Python Developer", "Python Engineer", "Backend Developer", "Data Engineer"],
            "PHP": ["PHP Developer", "Web Developer", "Full Stack Developer", "Backend PHP Developer"],
            "AI": ["AI Engineer", "Machine Learning Engineer", "Data Scientist", "AI/ML Developer"],
            "React": ["React Developer", "Frontend Developer", "React Engineer", "Web Developer"],
            "Frontend": ["Frontend Developer", "UI Developer", "Web Developer", "JavaScript Developer"],
            "Backend": ["Backend Developer", "Software Engineer", "API Developer", "Server-Side Developer"],
            "Full Stack": ["Full Stack Developer", "Software Engineer", "Web Developer", "Application Developer"],
            "DevOps": ["DevOps Engineer", "Site Reliability Engineer", "Cloud Engineer", "Infrastructure Engineer"],
            "Mobile": ["Mobile Developer", "iOS Developer", "Android Developer", "React Native Developer"],
        }
        
        job_titles = tech_titles.get(detected_tech, ["Software Developer", "Software Engineer", "Application Developer", "Technical Specialist"])
        
        version_chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"auto_versions_{resume_id}_{uuid.uuid4().hex[:8]}",
            system_message="""You are an expert resume writer. Create variations optimized for different job titles."""
        ).with_model("openai", "gpt-5.2")
        
        for title in job_titles[:4]:
            version_prompt = f"""Create a resume version optimized for: {title}

BASE RESUME:
{results["master_resume"]}

TECHNOLOGY: {detected_tech or 'Software Development'}

Modify for "{title}":
1. Update PROFESSIONAL SUMMARY for this title
2. Emphasize relevant skills
3. Adjust bullet points for this role
4. Use keywords common for "{title}"

Return ONLY the modified resume in plain text."""

            version_message = UserMessage(text=version_prompt)
            version_content = await version_chat.send_message(version_message)
            
            results["title_versions"].append({
                "name": title,
                "content": version_content,
                "job_title": title
            })
        
        logger.info(f"Generated {len(results['title_versions'])} versions for {resume_id}")
        
        # Update resume document with all results
        await db.resumes.update_one(
            {"resume_id": resume_id},
            {"$set": {
                "analysis": results["analysis"],
                "analyzed_at": datetime.now(timezone.utc).isoformat(),
                "master_resume": results["master_resume"],
                "master_created_at": datetime.now(timezone.utc).isoformat(),
                "title_versions": results["title_versions"],
                "versions_created_at": datetime.now(timezone.utc).isoformat(),
                "auto_processed": True,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        
    except Exception as e:
        logger.error(f"Auto-analysis error for {resume_id}: {str(e)}")
        results["error"] = str(e)
    
    return results

@api_router.post("/resumes/upload")
async def upload_resume(
    request: Request,
    file: UploadFile = File(...)
):
    user = await get_current_user(request)
    
    # Check resume limit (max 5 resumes per user)
    resume_count = await db.resumes.count_documents({"user_id": user["user_id"]})
    if resume_count >= 5:
        raise HTTPException(
            status_code=400, 
            detail="Maximum 5 resumes allowed. Please delete an existing resume to upload a new one."
        )
    
    content = await file.read()
    file_extension = file.filename.split('.')[-1].lower()
    
    # Extract text from file
    text_content = ""
    if file_extension == 'pdf':
        try:
            pdf_reader = PdfReader(BytesIO(content))
            for page in pdf_reader.pages:
                text_content += page.extract_text() or ""
        except Exception as e:
            text_content = f"Error extracting PDF: {str(e)}"
    elif file_extension in ['doc', 'docx']:
        try:
            doc = Document(BytesIO(content))
            text_content = "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            text_content = f"Error extracting Word doc: {str(e)}"
    else:
        text_content = content.decode('utf-8', errors='ignore')
    
    resume_id = f"resume_{uuid.uuid4().hex[:12]}"
    
    # Check if this is the first resume (make it primary)
    is_first_resume = resume_count == 0
    
    resume_doc = {
        "resume_id": resume_id,
        "user_id": user["user_id"],
        "file_name": file.filename,
        "file_type": file_extension,
        "original_content": text_content,
        "file_data": base64.b64encode(content).decode('utf-8'),
        "tailored_content": None,
        "auto_processed": False,
        "is_primary": is_first_resume,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.resumes.insert_one(resume_doc)
    
    # Automatically analyze, extract profile, create master, and generate versions
    logger.info(f"Starting auto-analysis for resume {resume_id}")
    auto_results = await auto_analyze_resume(resume_id, text_content, user["user_id"])
    
    # Get the updated user profile to return to frontend
    updated_user = await db.users.find_one({"user_id": user["user_id"]}, {"_id": 0, "password": 0})
    
    return {
        "resume_id": resume_id,
        "file_name": file.filename,
        "content_preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
        "message": "Resume uploaded and analyzed successfully. Profile updated from resume.",
        "analysis": auto_results.get("analysis"),
        "master_resume": auto_results.get("master_resume"),
        "title_versions": auto_results.get("title_versions", []),
        "extracted_profile": auto_results.get("extracted_profile"),
        "updated_user": updated_user,
        "auto_processed": True,
        "profile_updated": auto_results.get("extracted_profile") is not None
    }

@api_router.get("/resumes")
async def get_resumes(request: Request):
    user = await get_current_user(request)
    resumes = await db.resumes.find(
        {"user_id": user["user_id"]},
        {"_id": 0, "file_data": 0}
    ).to_list(100)
    return resumes

@api_router.get("/resumes/{resume_id}")
async def get_resume(resume_id: str, request: Request):
    user = await get_current_user(request)
    resume = await db.resumes.find_one(
        {"resume_id": resume_id, "user_id": user["user_id"]},
        {"_id": 0}
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return resume

@api_router.delete("/resumes/{resume_id}")
async def delete_resume(resume_id: str, request: Request):
    user = await get_current_user(request)
    result = await db.resumes.delete_one(
        {"resume_id": resume_id, "user_id": user["user_id"]}
    )
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Resume not found")
    return {"message": "Resume deleted successfully"}

@api_router.put("/resumes/{resume_id}/set-primary")
async def set_primary_resume(resume_id: str, request: Request):
    """Set a resume as the primary resume for job applications"""
    user = await get_current_user(request)
    
    # Check if resume exists
    resume = await db.resumes.find_one(
        {"resume_id": resume_id, "user_id": user["user_id"]}
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Remove primary flag from all other resumes
    await db.resumes.update_many(
        {"user_id": user["user_id"]},
        {"$set": {"is_primary": False}}
    )
    
    # Set this resume as primary
    await db.resumes.update_one(
        {"resume_id": resume_id},
        {"$set": {"is_primary": True}}
    )
    
    return {"message": "Resume set as primary successfully", "resume_id": resume_id}

@api_router.post("/resumes/tailor")
async def tailor_resume(data: TailorResumeRequest, request: Request):
    user = await get_current_user(request)
    
    resume = await db.resumes.find_one(
        {"resume_id": data.resume_id, "user_id": user["user_id"]},
        {"_id": 0}
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Enhanced ATS-optimized system message
    system_message = """You are an expert resume writer, ATS optimization specialist, and career consultant.
Your expertise includes:
- Making resumes ATS (Applicant Tracking System) friendly
- Identifying and incorporating relevant keywords from job descriptions
- Formatting content for both human readers and automated parsing systems
- Highlighting quantifiable achievements and metrics
- Using industry-standard section headers (SUMMARY, EXPERIENCE, SKILLS, EDUCATION)

ATS Optimization Rules:
1. Use standard section headers that ATS systems recognize
2. Include relevant keywords naturally throughout the resume
3. Avoid tables, graphics, headers/footers that ATS cannot parse
4. Use standard fonts and formatting
5. Include both spelled-out terms and acronyms (e.g., "Artificial Intelligence (AI)")
6. Place most important keywords in the top third of the resume
7. Use reverse chronological order for experience
8. Include job titles that match the target position terminology"""

    # Use AI to tailor resume with ATS optimization
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"tailor_{data.resume_id}_{uuid.uuid4().hex[:8]}",
        system_message=system_message
    ).with_model("openai", "gpt-5.2")
    
    # Extract keywords from job description
    keywords_prompt = f"""Analyze this job description and extract the top 15-20 most important keywords and phrases that an ATS would look for:

Job Title: {data.job_title}
Company: {data.company_name or 'Not specified'}
Technologies: {', '.join(data.technologies) if data.technologies else 'Not specified'}

Job Description:
{data.job_description}

Return ONLY a comma-separated list of keywords, nothing else."""

    keywords_message = UserMessage(text=keywords_prompt)
    extracted_keywords = await chat.send_message(keywords_message)
    
    # Use custom prompt if provided, otherwise use default
    if data.custom_prompt:
        prompt = f"""{data.custom_prompt}

ORIGINAL RESUME:
{resume['original_content']}

TARGET KEYWORDS (incorporate these naturally):
{extracted_keywords}

Return ONLY the tailored resume content, formatted as plain text with clear section headers."""
    else:
        # Main tailoring prompt
        prompt = f"""Tailor the following resume for the position of {data.job_title} at {data.company_name or 'the company'}.

TARGET KEYWORDS TO INCORPORATE (from job description):
{extracted_keywords}

REQUIRED TECHNOLOGIES TO HIGHLIGHT:
{', '.join(data.technologies) if data.technologies else 'As mentioned in the job description'}

JOB DESCRIPTION:
{data.job_description}

ORIGINAL RESUME:
{resume['original_content']}

CREATE AN ATS-OPTIMIZED RESUME that:
1. Uses standard ATS-friendly section headers: PROFESSIONAL SUMMARY, SKILLS, EXPERIENCE, EDUCATION, CERTIFICATIONS
2. Incorporates the target keywords naturally (aim for 70%+ keyword match)
3. Starts with a powerful PROFESSIONAL SUMMARY (3-4 lines) with key skills
4. Lists a SKILLS section with relevant technical and soft skills
5. Uses bullet points for achievements, starting with strong action verbs
6. Includes quantifiable metrics where possible (percentages, numbers, dollar amounts)
7. Maintains reverse chronological order
8. Uses clear, ATS-parseable formatting

Return ONLY the tailored resume content, formatted as plain text with clear section headers."""

    message = UserMessage(text=prompt)
    tailored_content = await chat.send_message(message)
    
    versions = []
    
    # Generate additional versions if requested
    if data.generate_versions:
        # Version 2: More technical focus
        version2_prompt = f"""Create an alternative version of this tailored resume with MORE TECHNICAL FOCUS.
        
Original tailored resume:
{tailored_content}

For this version:
1. Lead with technical skills and certifications
2. Emphasize technical projects and implementations
3. Include more technical keywords and acronyms
4. Focus on tools, technologies, and methodologies used
5. Highlight technical problem-solving achievements

Keep it ATS-friendly. Return ONLY the resume content."""

        version2_message = UserMessage(text=version2_prompt)
        version2_content = await chat.send_message(version2_message)
        
        # Version 3: Leadership/Impact focus
        version3_prompt = f"""Create an alternative version of this tailored resume with LEADERSHIP & IMPACT FOCUS.
        
Original tailored resume:
{tailored_content}

For this version:
1. Emphasize leadership roles and team management
2. Highlight business impact and ROI of projects
3. Focus on stakeholder management and communication
4. Include mentoring and cross-functional collaboration
5. Emphasize strategic thinking and decision-making

Keep it ATS-friendly. Return ONLY the resume content."""

        version3_message = UserMessage(text=version3_prompt)
        version3_content = await chat.send_message(version3_message)
        
        versions = [
            {"name": "Standard ATS-Optimized", "content": tailored_content},
            {"name": "Technical Focus", "content": version2_content},
            {"name": "Leadership Focus", "content": version3_content}
        ]
    
    # Update resume with tailored content
    update_data = {
        "tailored_content": tailored_content,
        "target_job_title": data.job_title,
        "target_job_description": data.job_description,
        "target_technologies": data.technologies,
        "extracted_keywords": extracted_keywords,
        "ats_optimized": True,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    if versions:
        update_data["versions"] = versions
    
    await db.resumes.update_one(
        {"resume_id": data.resume_id},
        {"$set": update_data}
    )
    
    response_data = {
        "resume_id": data.resume_id,
        "tailored_content": tailored_content,
        "keywords": extracted_keywords,
        "ats_optimized": True,
        "message": "Resume tailored successfully with ATS optimization"
    }
    
    if versions:
        response_data["versions"] = versions
    
    return response_data

@api_router.post("/resumes/{resume_id}/optimize")
async def optimize_resume_ats(resume_id: str, data: OptimizeResumeRequest, request: Request):
    """Make an uploaded resume ATS-friendly with keyword extraction and optional version generation"""
    user = await get_current_user(request)
    
    resume = await db.resumes.find_one(
        {"resume_id": resume_id, "user_id": user["user_id"]},
        {"_id": 0}
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    original_content = resume.get('original_content', '')
    if not original_content:
        raise HTTPException(status_code=400, detail="Resume has no content to optimize")
    
    # System message for ATS optimization
    system_message = """You are an expert ATS (Applicant Tracking System) optimization specialist and professional resume writer.
Your task is to transform resumes to be ATS-friendly while maintaining authenticity and readability.

ATS Optimization Guidelines:
1. Use standard, recognizable section headers: PROFESSIONAL SUMMARY, SKILLS, EXPERIENCE, EDUCATION, CERTIFICATIONS
2. Remove fancy formatting, tables, columns, graphics that ATS cannot parse
3. Use standard bullet points (•) for lists
4. Include both acronyms and full terms (e.g., "Artificial Intelligence (AI)")
5. Place most important skills and keywords in the top third
6. Use reverse chronological order for experience
7. Include quantifiable achievements with metrics (%, $, numbers)
8. Use industry-standard job titles
9. Ensure consistent date formatting (Month Year - Month Year)
10. Remove headers/footers that might confuse ATS"""

    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"optimize_{resume_id}_{uuid.uuid4().hex[:8]}",
        system_message=system_message
    ).with_model("openai", "gpt-5.2")
    
    # Extract keywords from the resume
    keywords_prompt = f"""Analyze this resume and extract the top 20 most important professional keywords that should be highlighted for ATS systems. Include:
- Technical skills and tools
- Industry-specific terms
- Certifications and qualifications
- Soft skills mentioned
- Action verbs used

Resume:
{original_content}

Return ONLY a comma-separated list of keywords, nothing else."""

    keywords_message = UserMessage(text=keywords_prompt)
    extracted_keywords = await chat.send_message(keywords_message)
    
    # Determine target role
    target_role = data.target_role if data.target_role else "a professional role matching their experience"
    
    # Main ATS optimization prompt
    optimize_prompt = f"""Transform this resume to be fully ATS-optimized for {target_role}.

CURRENT RESUME:
{original_content}

EXTRACTED KEYWORDS TO EMPHASIZE:
{extracted_keywords}

CREATE AN ATS-OPTIMIZED VERSION that:
1. Starts with a powerful PROFESSIONAL SUMMARY (3-4 lines) highlighting key qualifications
2. Includes a comprehensive SKILLS section organized by category (Technical Skills, Tools, Soft Skills)
3. Reformats EXPERIENCE section with:
   - Clear job titles and company names
   - Date ranges in consistent format
   - Bullet points starting with strong action verbs
   - Quantifiable achievements where possible
4. Includes EDUCATION with degrees, institutions, and graduation dates
5. Adds CERTIFICATIONS section if applicable
6. Uses clean, ATS-parseable formatting throughout
7. Naturally incorporates the extracted keywords

Return ONLY the optimized resume content in plain text format with clear section headers."""

    optimize_message = UserMessage(text=optimize_prompt)
    optimized_content = await chat.send_message(optimize_message)
    
    versions = []
    
    # Generate additional versions if requested
    if data.generate_versions:
        # Version 2: Technical/Skills-focused
        version2_prompt = f"""Create an alternative TECHNICAL FOCUS version of this resume.

Optimized resume:
{optimized_content}

For this version:
1. Lead with a TECHNICAL SKILLS section at the top (after contact info)
2. Emphasize technical projects, implementations, and tools
3. Include more technical acronyms and certifications
4. Highlight problem-solving and technical achievements
5. Focus on technologies, frameworks, and methodologies

Keep it ATS-friendly. Return ONLY the resume content."""

        version2_message = UserMessage(text=version2_prompt)
        version2_content = await chat.send_message(version2_message)
        
        # Version 3: Experience/Leadership-focused
        version3_prompt = f"""Create an alternative LEADERSHIP & EXPERIENCE FOCUS version of this resume.

Optimized resume:
{optimized_content}

For this version:
1. Lead with the PROFESSIONAL SUMMARY emphasizing leadership
2. Highlight team management and mentoring experience
3. Emphasize business impact, ROI, and cost savings
4. Focus on stakeholder management and communication
5. Include project management and strategic initiatives

Keep it ATS-friendly. Return ONLY the resume content."""

        version3_message = UserMessage(text=version3_prompt)
        version3_content = await chat.send_message(version3_message)
        
        versions = [
            {"name": "Standard ATS-Optimized", "content": optimized_content},
            {"name": "Technical Focus", "content": version2_content},
            {"name": "Leadership Focus", "content": version3_content}
        ]
    
    # Update resume in database
    update_data = {
        "ats_optimized": True,
        "ats_optimized_content": optimized_content,
        "extracted_keywords": extracted_keywords,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    if versions:
        update_data["versions"] = versions
    
    await db.resumes.update_one(
        {"resume_id": resume_id},
        {"$set": update_data}
    )
    
    response_data = {
        "resume_id": resume_id,
        "optimized_content": optimized_content,
        "keywords": extracted_keywords,
        "ats_optimized": True,
        "message": "Resume optimized for ATS successfully"
    }
    
    if versions:
        response_data["versions"] = versions
    
    return response_data

@api_router.post("/resumes/{resume_id}/generate-word")
async def generate_word_resume(resume_id: str, request: Request):
    """Generate a Word document from the tailored resume"""
    user = await get_current_user(request)
    
    # Try to get custom content from request body
    body = {}
    try:
        body = await request.json()
    except:
        pass
    
    custom_content = body.get("content")
    version = body.get("version", "default")
    
    resume = await db.resumes.find_one(
        {"resume_id": resume_id, "user_id": user["user_id"]},
        {"_id": 0}
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Get content - priority: custom_content > version > tailored > original
    if custom_content:
        content = custom_content
    elif version != "default" and resume.get("versions"):
        content = None
        for v in resume["versions"]:
            if v["name"] == version:
                content = v["content"]
                break
        if not content:
            content = resume.get("tailored_content") or resume.get("original_content", "")
    else:
        content = resume.get("tailored_content") or resume.get("original_content", "")
    
    # Create Word document
    doc = Document()
    
    # Set narrow margins for more content
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Add name as title
    name = user.get("name", "Candidate")
    title = doc.add_heading(name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.bold = True
    
    # Add contact info
    contact_parts = []
    if user.get('email'):
        contact_parts.append(user['email'])
    if user.get('phone'):
        contact_parts.append(user['phone'])
    if user.get('location'):
        contact_parts.append(user['location'])
    if user.get('linkedin_profile'):
        contact_parts.append(user['linkedin_profile'])
    
    if contact_parts:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)
    
    # Add horizontal line
    doc.add_paragraph("_" * 80)
    
    # Parse and add content with proper formatting
    lines = content.split('\n')
    current_section = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if it's a section header (ALL CAPS or common headers)
        section_headers = ['PROFESSIONAL SUMMARY', 'SUMMARY', 'SKILLS', 'TECHNICAL SKILLS', 
                         'EXPERIENCE', 'WORK EXPERIENCE', 'EDUCATION', 'CERTIFICATIONS',
                         'PROJECTS', 'ACHIEVEMENTS', 'AWARDS', 'LANGUAGES']
        
        is_header = line.upper() in section_headers or (line.isupper() and len(line) < 50)
        
        if is_header:
            # Add section header
            heading = doc.add_heading(line.title(), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading.runs:
                run.font.size = Pt(12)
                run.font.bold = True
            current_section = line.upper()
        elif line.startswith(('•', '-', '*', '○')):
            # Bullet point
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(line.lstrip('•-*○ '))
            run.font.size = Pt(10)
        elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            # Numbered list
            para = doc.add_paragraph(style='List Number')
            run = para.add_run(line[2:].strip())
            run.font.size = Pt(10)
        else:
            # Regular paragraph
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(10)
    
    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    # Generate filename
    job_title = resume.get("target_job_title", "").replace(" ", "_")[:30]
    filename = f"{name.replace(' ', '_')}_Resume_{job_title}.docx"
    
    return StreamingResponse(
        doc_io,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@api_router.get("/resumes/{resume_id}/download/{format}")
async def download_resume(resume_id: str, format: str, request: Request):
    user = await get_current_user(request)
    
    resume = await db.resumes.find_one(
        {"resume_id": resume_id, "user_id": user["user_id"]},
        {"_id": 0}
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    content = resume.get("tailored_content") or resume.get("original_content", "")
    
    if format == "docx":
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(user.get("name", "Resume"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact info
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(f"{user.get('email', '')} | {user.get('phone', '')} | {user.get('location', '')}")
        
        # Add content
        for paragraph in content.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=resume_{resume_id}.docx"}
        )
    
    elif format == "pdf":
        # Create PDF
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        # Add content
        y = height - inch
        c.setFont("Helvetica-Bold", 16)
        c.drawCentredString(width/2, y, user.get("name", "Resume"))
        
        y -= 30
        c.setFont("Helvetica", 10)
        c.drawCentredString(width/2, y, f"{user.get('email', '')} | {user.get('phone', '')} | {user.get('location', '')}")
        
        y -= 40
        c.setFont("Helvetica", 11)
        
        for line in content.split('\n'):
            if y < inch:
                c.showPage()
                y = height - inch
                c.setFont("Helvetica", 11)
            
            if line.strip():
                # Handle long lines
                words = line.split()
                current_line = ""
                for word in words:
                    test_line = current_line + " " + word if current_line else word
                    if c.stringWidth(test_line, "Helvetica", 11) < width - 2*inch:
                        current_line = test_line
                    else:
                        c.drawString(inch, y, current_line)
                        y -= 15
                        current_line = word
                if current_line:
                    c.drawString(inch, y, current_line)
                    y -= 15
            else:
                y -= 10
        
        c.save()
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=resume_{resume_id}.pdf"}
        )
    
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use 'docx' or 'pdf'")

# ============ RESUME ANALYSIS & ENHANCEMENT ============

@api_router.post("/resumes/{resume_id}/analyze")
async def analyze_resume(resume_id: str, request: Request):
    """Analyze resume and provide score, missing info, and improvement suggestions"""
    user = await get_current_user(request)
    
    resume = await db.resumes.find_one(
        {"resume_id": resume_id, "user_id": user["user_id"]},
        {"_id": 0}
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    original_content = resume.get('original_content', '')
    if not original_content:
        raise HTTPException(status_code=400, detail="Resume has no content to analyze")
    
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"analyze_{resume_id}_{uuid.uuid4().hex[:8]}",
        system_message="""You are an expert resume analyst and career consultant. Analyze resumes thoroughly and provide actionable feedback."""
    ).with_model("openai", "gpt-5.2")
    
    analysis_prompt = f"""Analyze this resume and provide a comprehensive assessment.

RESUME:
{original_content}

Provide your analysis in the following JSON format (return ONLY valid JSON, no markdown):
{{
    "score": <number 0-100>,
    "grade": "<A/B/C/D/F>",
    "summary": "<2-3 sentence overall assessment>",
    "missing_info": {{
        "phone": <true if missing, false if present>,
        "email": <true if missing, false if present>,
        "address_location": <true if missing, false if present>,
        "linkedin": <true if missing, false if present>,
        "professional_summary": <true if missing, false if present>,
        "skills_section": <true if missing, false if present>,
        "education": <true if missing, false if present>,
        "work_experience": <true if missing, false if present>,
        "quantifiable_achievements": <true if missing/weak, false if strong>
    }},
    "strengths": ["<strength 1>", "<strength 2>", "<strength 3>"],
    "weaknesses": ["<weakness 1>", "<weakness 2>", "<weakness 3>"],
    "improvement_suggestions": ["<suggestion 1>", "<suggestion 2>", "<suggestion 3>", "<suggestion 4>"],
    "ats_compatibility": {{
        "score": <number 0-100>,
        "issues": ["<issue 1>", "<issue 2>"]
    }},
    "detected_skills": ["<skill1>", "<skill2>", ...],
    "detected_job_titles": ["<title1>", "<title2>", ...],
    "experience_level": "<Entry/Mid/Senior/Executive>"
}}"""

    analysis_message = UserMessage(text=analysis_prompt)
    analysis_response = await chat.send_message(analysis_message)
    
    # Parse the JSON response
    try:
        import json
        # Clean up the response - remove any markdown formatting
        response_text = analysis_response.strip()
        if response_text.startswith('```'):
            response_text = response_text.split('\n', 1)[1]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
        analysis_data = json.loads(response_text)
    except:
        # Fallback if JSON parsing fails
        analysis_data = {
            "score": 50,
            "grade": "C",
            "summary": "Unable to parse detailed analysis. Please try again.",
            "missing_info": {},
            "strengths": [],
            "weaknesses": [],
            "improvement_suggestions": ["Upload a clearer resume for better analysis"],
            "ats_compatibility": {"score": 50, "issues": []},
            "detected_skills": [],
            "detected_job_titles": [],
            "experience_level": "Unknown"
        }
    
    # Store analysis in database
    await db.resumes.update_one(
        {"resume_id": resume_id},
        {"$set": {
            "analysis": analysis_data,
            "analyzed_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {
        "resume_id": resume_id,
        "analysis": analysis_data
    }

@api_router.post("/resumes/{resume_id}/create-master")
async def create_master_resume(resume_id: str, request: Request):
    """Create a polished master resume without a specific job description"""
    user = await get_current_user(request)
    
    resume = await db.resumes.find_one(
        {"resume_id": resume_id, "user_id": user["user_id"]},
        {"_id": 0}
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    original_content = resume.get('original_content', '')
    if not original_content:
        raise HTTPException(status_code=400, detail="Resume has no content")
    
    # Get user profile for additional context
    user_profile = await db.users.find_one({"user_id": user["user_id"]}, {"_id": 0})
    primary_tech = user_profile.get("primary_technology", "") if user_profile else ""
    sub_techs = user_profile.get("sub_technologies", []) if user_profile else []
    
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"master_{resume_id}_{uuid.uuid4().hex[:8]}",
        system_message="""You are an expert resume writer and career consultant. 
Your task is to create a polished, professional master resume that can be used as a base for any job application.
Focus on fixing formatting issues, improving content quality, and making it ATS-friendly."""
    ).with_model("openai", "gpt-5.2")
    
    master_prompt = f"""Transform this resume into a polished MASTER RESUME that serves as a strong foundation for any job application.

ORIGINAL RESUME:
{original_content}

USER'S PRIMARY TECHNOLOGY: {primary_tech or 'Not specified'}
USER'S SKILLS: {', '.join(sub_techs) if sub_techs else 'Not specified'}

CREATE A MASTER RESUME that:
1. PROFESSIONAL SUMMARY: Write a compelling 3-4 line summary highlighting key expertise and value proposition
2. FIX ALL ISSUES: Grammar, spelling, formatting, inconsistencies
3. SKILLS SECTION: Organize skills into categories (Technical Skills, Tools & Platforms, Soft Skills)
4. EXPERIENCE: 
   - Rewrite bullet points with strong action verbs
   - Add quantifiable metrics where possible (even estimates like "Improved performance by ~20%")
   - Ensure consistent formatting (dates, company names, titles)
5. EDUCATION: Proper formatting with degrees, institutions, dates
6. FORMATTING: Use standard ATS-friendly sections:
   - PROFESSIONAL SUMMARY
   - SKILLS
   - PROFESSIONAL EXPERIENCE
   - EDUCATION
   - CERTIFICATIONS (if any)
7. CONTACT INFO: Ensure placeholder for phone, email, location, LinkedIn

Make it professional, impactful, and ready to customize for specific jobs.
Return ONLY the master resume content in plain text format."""

    master_message = UserMessage(text=master_prompt)
    master_content = await chat.send_message(master_message)
    
    # Store master resume
    await db.resumes.update_one(
        {"resume_id": resume_id},
        {"$set": {
            "master_resume": master_content,
            "master_created_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {
        "resume_id": resume_id,
        "master_resume": master_content,
        "message": "Master resume created successfully"
    }

@api_router.post("/resumes/{resume_id}/generate-versions")
async def generate_resume_versions(resume_id: str, request: Request):
    """Generate 3-4 resume versions with different job title designations based on candidate's technology"""
    user = await get_current_user(request)
    
    # Get request body
    body = {}
    try:
        body = await request.json()
    except:
        pass
    
    resume = await db.resumes.find_one(
        {"resume_id": resume_id, "user_id": user["user_id"]},
        {"_id": 0}
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Use master resume if available, otherwise use original
    base_content = resume.get('master_resume') or resume.get('tailored_content') or resume.get('original_content', '')
    if not base_content:
        raise HTTPException(status_code=400, detail="Resume has no content")
    
    # Get user's primary technology
    user_profile = await db.users.find_one({"user_id": user["user_id"]}, {"_id": 0})
    primary_tech = body.get("primary_technology") or (user_profile.get("primary_technology", "") if user_profile else "")
    
    # Define popular job titles for each technology
    tech_titles = {
        "Java": ["Java Developer", "Java Software Engineer", "Backend Developer", "Full Stack Java Developer"],
        "Python": ["Python Developer", "Python Engineer", "Backend Developer", "Data Engineer"],
        "PHP": ["PHP Developer", "Web Developer", "Full Stack Developer", "Backend PHP Developer"],
        "AI": ["AI Engineer", "Machine Learning Engineer", "Data Scientist", "AI/ML Developer"],
        "React": ["React Developer", "Frontend Developer", "React Engineer", "Web Developer"],
        "Frontend": ["Frontend Developer", "UI Developer", "Web Developer", "JavaScript Developer"],
        "Backend": ["Backend Developer", "Software Engineer", "API Developer", "Server-Side Developer"],
        "Full Stack": ["Full Stack Developer", "Software Engineer", "Web Developer", "Application Developer"],
        "DevOps": ["DevOps Engineer", "Site Reliability Engineer", "Cloud Engineer", "Infrastructure Engineer"],
        "Mobile": ["Mobile Developer", "iOS Developer", "Android Developer", "React Native Developer"],
    }
    
    # Get relevant titles or use generic ones
    job_titles = tech_titles.get(primary_tech, ["Software Developer", "Software Engineer", "Application Developer", "Technical Specialist"])
    
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"versions_{resume_id}_{uuid.uuid4().hex[:8]}",
        system_message="""You are an expert resume writer. Create variations of resumes optimized for different job titles while maintaining the candidate's actual experience and skills."""
    ).with_model("openai", "gpt-5.2")
    
    versions = []
    
    for i, title in enumerate(job_titles[:4]):  # Generate up to 4 versions
        version_prompt = f"""Create a resume version optimized for the job title: {title}

BASE RESUME:
{base_content}

PRIMARY TECHNOLOGY: {primary_tech or 'Software Development'}

Modify this resume to be optimized for a "{title}" position:
1. Update the PROFESSIONAL SUMMARY to highlight relevant skills for this title
2. Reorder and emphasize skills most relevant to this role
3. Adjust experience bullet points to emphasize relevant work
4. Use keywords and terminology common for "{title}" positions
5. Keep all factual information accurate - only adjust emphasis and presentation

Return ONLY the modified resume content in plain text format."""

        version_message = UserMessage(text=version_prompt)
        version_content = await chat.send_message(version_message)
        
        versions.append({
            "name": title,
            "content": version_content,
            "job_title": title
        })
    
    # Store versions in database
    await db.resumes.update_one(
        {"resume_id": resume_id},
        {"$set": {
            "title_versions": versions,
            "versions_created_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {
        "resume_id": resume_id,
        "versions": versions,
        "primary_technology": primary_tech,
        "message": f"Generated {len(versions)} resume versions for different job titles"
    }

# ============ COVER LETTER ============

@api_router.post("/cover-letter/generate")
async def generate_cover_letter(data: GenerateCoverLetterRequest, request: Request):
    user = await get_current_user(request)
    
    resume = await db.resumes.find_one(
        {"resume_id": data.resume_id, "user_id": user["user_id"]},
        {"_id": 0}
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"cover_{data.resume_id}_{uuid.uuid4().hex[:8]}",
        system_message="""You are an expert at writing compelling cover letters.
        Create personalized, professional cover letters that highlight relevant experience
        and show enthusiasm for the role and company."""
    ).with_model("openai", "gpt-5.2")
    
    prompt = f"""Write a professional cover letter for the following position:

Position: {data.job_title}
Company: {data.company_name}
Job Description: {data.job_description}

Candidate Name: {user.get('name', 'Candidate')}
Candidate Email: {user.get('email', '')}

Resume Content:
{resume.get('tailored_content') or resume.get('original_content', '')}

Please write a compelling cover letter that:
1. Opens with a strong hook
2. Highlights 2-3 most relevant experiences
3. Shows knowledge of the company (if mentioned in job description)
4. Expresses genuine interest in the role
5. Ends with a clear call to action

Keep it to one page (about 300-400 words)."""

    message = UserMessage(text=prompt)
    cover_letter = await chat.send_message(message)
    
    return {
        "cover_letter": cover_letter,
        "job_title": data.job_title,
        "company_name": data.company_name
    }

# ============ JOB PORTALS ============

@api_router.post("/job-portals")
async def create_job_portal(data: JobPortalCreate, request: Request):
    await get_admin_user(request)
    
    portal_id = f"portal_{uuid.uuid4().hex[:12]}"
    portal_doc = {
        "portal_id": portal_id,
        "name": data.name,
        "url": data.url,
        "technology": data.technology,
        "description": data.description,
        "is_active": True,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.job_portals.insert_one(portal_doc)
    return {"portal_id": portal_id, "message": "Job portal created successfully"}

@api_router.get("/job-portals")
async def get_job_portals(request: Request, technology: Optional[str] = None):
    await get_current_user(request)
    
    query = {"is_active": True}
    if technology:
        query["technology"] = technology
    
    portals = await db.job_portals.find(query, {"_id": 0}).to_list(100)
    return portals

@api_router.put("/job-portals/{portal_id}")
async def update_job_portal(portal_id: str, data: JobPortalCreate, request: Request):
    await get_admin_user(request)
    
    result = await db.job_portals.update_one(
        {"portal_id": portal_id},
        {"$set": {
            "name": data.name,
            "url": data.url,
            "technology": data.technology,
            "description": data.description,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Portal not found")
    
    return {"message": "Portal updated successfully"}

@api_router.delete("/job-portals/{portal_id}")
async def delete_job_portal(portal_id: str, request: Request):
    await get_admin_user(request)
    
    result = await db.job_portals.update_one(
        {"portal_id": portal_id},
        {"$set": {"is_active": False}}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Portal not found")
    
    return {"message": "Portal deleted successfully"}

# ============ JOB APPLICATIONS ============

@api_router.post("/applications")
async def create_application(data: JobApplicationCreate, request: Request):
    user = await get_current_user(request)
    
    application_id = f"app_{uuid.uuid4().hex[:12]}"
    
    # Get the original resume info
    resume_info = None
    if data.resume_id:
        resume = await db.resumes.find_one(
            {"resume_id": data.resume_id, "user_id": user["user_id"]},
            {"_id": 0, "file_name": 1, "original_content": 1}
        )
        if resume:
            resume_info = {
                "file_name": resume.get("file_name"),
                "original_content": resume.get("original_content", "")[:500]  # Preview
            }
    
    application_doc = {
        "application_id": application_id,
        "user_id": user["user_id"],
        "job_portal_id": data.job_portal_id,
        "job_title": data.job_title,
        "job_description": data.job_description,
        "company_name": data.company_name,
        "resume_id": data.resume_id,
        "cover_letter": data.cover_letter,
        "tailored_content": data.tailored_content,  # Save the tailored resume
        "job_source": data.job_source,
        "apply_link": data.apply_link,
        "resume_info": resume_info,
        "status": "applied",
        "submission_screenshot": None,  # Will be updated when screenshot is uploaded
        "applied_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.applications.insert_one(application_doc)
    
    return {"application_id": application_id, "message": "Application submitted successfully"}

@api_router.post("/applications/{application_id}/screenshot")
async def upload_submission_screenshot(
    application_id: str,
    request: Request,
    file: UploadFile = File(...)
):
    """Upload a screenshot of the job submission page"""
    user = await get_current_user(request)
    
    # Verify application belongs to user
    application = await db.applications.find_one({
        "application_id": application_id,
        "user_id": user["user_id"]
    })
    
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
    
    # Read and store screenshot
    content = await file.read()
    screenshot_data = base64.b64encode(content).decode('utf-8')
    
    # Update application with screenshot
    await db.applications.update_one(
        {"application_id": application_id},
        {"$set": {
            "submission_screenshot": screenshot_data,
            "screenshot_filename": file.filename,
            "screenshot_uploaded_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {"message": "Screenshot uploaded successfully"}

@api_router.get("/applications")
async def get_applications(request: Request, status: Optional[str] = None):
    user = await get_current_user(request)
    
    query = {"user_id": user["user_id"]}
    if status:
        query["status"] = status
    
    applications = await db.applications.find(query, {"_id": 0}).to_list(100)
    return applications

@api_router.put("/applications/{application_id}/status")
async def update_application_status(
    application_id: str,
    status: str,
    request: Request
):
    user = await get_current_user(request)
    
    valid_statuses = ["applied", "screening", "interview_scheduled", "interviewed", "offer", "rejected", "withdrawn"]
    if status not in valid_statuses:
        raise HTTPException(status_code=400, detail=f"Invalid status. Must be one of: {valid_statuses}")
    
    result = await db.applications.update_one(
        {"application_id": application_id, "user_id": user["user_id"]},
        {"$set": {"status": status, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Application not found")
    
    return {"message": "Application status updated"}

# ============ EMAIL COMMUNICATION ============

@api_router.post("/emails")
async def create_email(data: EmailCreate, request: Request):
    user = await get_current_user(request)
    
    email_id = f"email_{uuid.uuid4().hex[:12]}"
    email_doc = {
        "email_id": email_id,
        "user_id": user["user_id"],
        "application_id": data.application_id,
        "subject": data.subject,
        "content": data.content,
        "email_type": data.email_type,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.emails.insert_one(email_doc)
    return {"email_id": email_id, "message": "Email recorded successfully"}

@api_router.get("/emails")
async def get_emails(request: Request, application_id: Optional[str] = None):
    user = await get_current_user(request)
    
    query = {"user_id": user["user_id"]}
    if application_id:
        query["application_id"] = application_id
    
    emails = await db.emails.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    return emails

@api_router.post("/emails/generate-reply")
async def generate_email_reply(data: EmailReplyRequest, request: Request):
    await get_current_user(request)
    
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"email_{uuid.uuid4().hex[:8]}",
        system_message=f"""You are a professional email assistant.
        Write {data.tone} email replies that are clear, concise, and appropriate for job applications."""
    ).with_model("openai", "gpt-5.2")
    
    prompt = f"""Generate a professional reply to the following email:

Original Email:
{data.original_email}

Context/Instructions:
{data.context}

Please write a {data.tone} response that:
1. Addresses all points in the original email
2. Is professional and appropriate for job application correspondence
3. Is concise but complete
4. Includes appropriate greeting and sign-off"""

    message = UserMessage(text=prompt)
    reply = await chat.send_message(message)
    
    return {"generated_reply": reply}

# ============ REPORTS ============

@api_router.get("/reports/candidate")
async def get_candidate_report(request: Request):
    user = await get_current_user(request)
    user_id = user["user_id"]
    
    # Get application stats
    total_applications = await db.applications.count_documents({"user_id": user_id})
    
    # Status breakdown
    pipeline = [
        {"$match": {"user_id": user_id}},
        {"$group": {"_id": "$status", "count": {"$sum": 1}}}
    ]
    status_breakdown = await db.applications.aggregate(pipeline).to_list(100)
    
    # Recent applications
    recent = await db.applications.find(
        {"user_id": user_id},
        {"_id": 0}
    ).sort("applied_at", -1).limit(5).to_list(5)
    
    # Resume count
    resume_count = await db.resumes.count_documents({"user_id": user_id})
    
    # Email count
    email_count = await db.emails.count_documents({"user_id": user_id})
    
    return {
        "total_applications": total_applications,
        "status_breakdown": {item["_id"]: item["count"] for item in status_breakdown},
        "recent_applications": recent,
        "resume_count": resume_count,
        "email_count": email_count,
        "interviews_scheduled": sum(1 for item in status_breakdown if item["_id"] in ["interview_scheduled", "interviewed"]),
        "offers_received": sum(1 for item in status_breakdown if item["_id"] == "offer")
    }

@api_router.get("/reports/admin")
async def get_admin_report(request: Request):
    await get_admin_user(request)
    
    # Total candidates
    total_candidates = await db.users.count_documents({"role": "candidate"})
    
    # Total applications
    total_applications = await db.applications.count_documents({})
    
    # Applications by status
    pipeline = [
        {"$group": {"_id": "$status", "count": {"$sum": 1}}}
    ]
    status_breakdown = await db.applications.aggregate(pipeline).to_list(100)
    
    # Applications by technology
    tech_pipeline = [
        {"$lookup": {
            "from": "users",
            "localField": "user_id",
            "foreignField": "user_id",
            "as": "user"
        }},
        {"$unwind": "$user"},
        {"$group": {"_id": "$user.primary_technology", "count": {"$sum": 1}}}
    ]
    tech_breakdown = await db.applications.aggregate(tech_pipeline).to_list(100)
    
    # Recent registrations
    recent_users = await db.users.find(
        {"role": "candidate"},
        {"_id": 0, "password": 0}
    ).sort("created_at", -1).limit(10).to_list(10)
    
    # Active job portals
    active_portals = await db.job_portals.count_documents({"is_active": True})
    
    return {
        "total_candidates": total_candidates,
        "total_applications": total_applications,
        "status_breakdown": {item["_id"]: item["count"] for item in status_breakdown},
        "technology_breakdown": {item["_id"]: item["count"] for item in tech_breakdown if item["_id"]},
        "recent_registrations": recent_users,
        "active_portals": active_portals,
        "total_interviews": sum(item["count"] for item in status_breakdown if item["_id"] in ["interview_scheduled", "interviewed"]),
        "total_offers": sum(item["count"] for item in status_breakdown if item["_id"] == "offer")
    }

@api_router.get("/reports/admin/candidates")
async def get_all_candidates(request: Request, skip: int = 0, limit: int = 20):
    await get_admin_user(request)
    
    candidates = await db.users.find(
        {"role": "candidate"},
        {"_id": 0, "password": 0}
    ).skip(skip).limit(limit).to_list(limit)
    
    # Enrich with application stats
    for candidate in candidates:
        app_count = await db.applications.count_documents({"user_id": candidate["user_id"]})
        interview_count = await db.applications.count_documents({
            "user_id": candidate["user_id"],
            "status": {"$in": ["interview_scheduled", "interviewed"]}
        })
        candidate["total_applications"] = app_count
        candidate["total_interviews"] = interview_count
    
    total = await db.users.count_documents({"role": "candidate"})
    
    return {
        "candidates": candidates,
        "total": total,
        "skip": skip,
        "limit": limit
    }

# ============ ADMIN USER MANAGEMENT ============

@api_router.post("/admin/create")
async def create_admin(user_data: UserCreate, admin_secret: str):
    # Simple secret check - in production use proper admin creation flow
    if admin_secret != "admin-secret-key-2024":
        raise HTTPException(status_code=403, detail="Invalid admin secret")
    
    existing = await db.users.find_one({"email": user_data.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user_id = f"user_{uuid.uuid4().hex[:12]}"
    hashed_password = hash_password(user_data.password)
    
    user_doc = {
        "user_id": user_id,
        "email": user_data.email,
        "password": hashed_password,
        "name": user_data.name,
        "primary_technology": user_data.primary_technology,
        "sub_technologies": user_data.sub_technologies,
        "phone": user_data.phone,
        "location": user_data.location,
        "role": "admin",
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_doc)
    
    return {"user_id": user_id, "message": "Admin user created successfully"}

# ============ LIVE JOB SEARCH (JSearch API) ============

class LiveJobSearchRequest(BaseModel):
    query: Optional[str] = None
    location: Optional[str] = "United States"
    employment_type: Optional[str] = None
    page: int = 1
    num_pages: int = 1

@api_router.get("/live-jobs/search")
async def search_live_jobs(
    request: Request,
    query: Optional[str] = None,
    location: Optional[str] = "United States",
    employment_type: Optional[str] = None,
    source: Optional[str] = None,  # Filter by source: Indeed, Dice, RemoteOK, Arbeitnow
    page: int = 1
):
    """
    Search for live job listings from Indeed, Dice, RemoteOK, and other job boards.
    Uses web scraping to get real-time jobs without API dependencies.
    """
    user = await get_current_user(request)
    
    # Build search query from user's technologies if not provided
    if not query:
        technologies = []
        if user.get("primary_technology"):
            technologies.append(user["primary_technology"])
        if user.get("sub_technologies"):
            technologies.extend(user["sub_technologies"][:2])
        
        if technologies:
            query = technologies[0]  # Use primary technology for search
        else:
            query = "software developer"
    
    try:
        # Use web scraper to fetch jobs
        logger.info(f"Searching jobs for: {query} in {location}, source filter: {source}")
        
        scraped_jobs = await job_scraper.scrape_all_sources(
            query=query,
            location=location,
            limit_per_source=10
        )
        
        # Filter by source if specified
        if source and source.lower() != 'all':
            source_lower = source.lower()
            scraped_jobs = [
                job for job in scraped_jobs 
                if source_lower in (job.get('source', '') or '').lower()
            ]
        
        # Filter by employment type if specified
        if employment_type:
            employment_type_lower = employment_type.lower()
            scraped_jobs = [
                job for job in scraped_jobs 
                if employment_type_lower in (job.get('employment_type', '') or '').lower()
            ]
        
        return {
            "jobs": scraped_jobs,
            "total": len(scraped_jobs),
            "page": page,
            "query_used": query,
            "location": location,
            "source_filter": source,
            "sources": ["Indeed", "Dice", "RemoteOK", "Arbeitnow"],
            "data_source": "live_scraping"
        }
        
    except Exception as e:
        logger.error(f"Error searching jobs: {str(e)}")
        return {
            "jobs": [],
            "total": 0,
            "page": page,
            "query_used": query,
            "location": location,
            "error": str(e)
        }

@api_router.get("/live-jobs/recommendations")
async def get_job_recommendations(request: Request):
    """
    Get personalized job recommendations based on user's primary and sub technologies.
    Requires user to have primary_technology set.
    """
    user = await get_current_user(request)
    
    # Check if user has required profile fields
    if not user.get('primary_technology'):
        return {
            "recommendations": [],
            "message": "Please update your profile with Primary Technology to get personalized job recommendations.",
            "requires_profile_update": True,
            "missing_fields": ["primary_technology"]
        }
    
    primary_tech = user.get("primary_technology", "")
    sub_techs = user.get("sub_technologies", [])
    user_location = user.get("location", "United States")
    
    try:
        # Use web scraper to fetch real jobs from multiple sources
        logger.info(f"Scraping jobs for: {primary_tech} in {user_location}")
        
        # Scrape from all sources
        scraped_jobs = await job_scraper.scrape_all_sources(
            query=primary_tech,
            location=user_location,
            limit_per_source=10
        )
        
        if scraped_jobs:
            logger.info(f"Found {len(scraped_jobs)} jobs from web scraping")
            return {
                "recommendations": scraped_jobs[:20],  # Limit to 20 jobs
                "total": len(scraped_jobs),
                "based_on": {
                    "primary_technology": primary_tech,
                    "sub_technologies": sub_techs
                },
                "sources": ["Indeed", "Dice", "RemoteOK", "Arbeitnow"],
                "data_source": "live_scraping"
            }
        
        # If scraping returned nothing, try sub-technologies
        if sub_techs:
            for sub_tech in sub_techs[:2]:
                more_jobs = await job_scraper.scrape_all_sources(
                    query=sub_tech,
                    location=user_location,
                    limit_per_source=5
                )
                scraped_jobs.extend(more_jobs)
        
        if scraped_jobs:
            # Remove duplicates
            seen_ids = set()
            unique_jobs = []
            for job in scraped_jobs:
                if job['job_id'] not in seen_ids:
                    seen_ids.add(job['job_id'])
                    unique_jobs.append(job)
            
            return {
                "recommendations": unique_jobs[:20],
                "total": len(unique_jobs),
                "based_on": {
                    "primary_technology": primary_tech,
                    "sub_technologies": sub_techs
                },
                "sources": ["Indeed", "Dice", "RemoteOK", "Arbeitnow"],
                "data_source": "live_scraping"
            }
    
    except Exception as e:
        logger.error(f"Error scraping jobs: {str(e)}")
    
    # Fallback to sample jobs if scraping fails
    def get_sample_jobs(tech):
        sample_jobs = [
            {
                "job_id": f"sample_1_{tech.lower()}",
                "title": f"Senior {tech} Developer",
                "company": "TechCorp Inc.",
                "company_logo": "https://ui-avatars.com/api/?name=TC&background=6366f1&color=fff",
                "location": "San Francisco, CA",
                "employment_type": "Full-time",
                "salary_min": 150000,
                "salary_max": 200000,
                "description": f"We are looking for an experienced {tech} developer to join our team. You will be working on cutting-edge projects with modern technologies.",
                "apply_link": "https://example.com/apply",
                "posted_date": datetime.now(timezone.utc).isoformat(),
                "is_remote": True,
                "matched_technology": tech,
                "source": "Sample"
            },
            {
                "job_id": f"sample_2_{tech.lower()}",
                "title": f"{tech} Software Engineer",
                "company": "Innovation Labs",
                "company_logo": "https://ui-avatars.com/api/?name=IL&background=8b5cf6&color=fff",
                "location": "New York, NY",
                "employment_type": "Full-time",
                "salary_min": 130000,
                "salary_max": 180000,
                "description": f"Join our dynamic team as a {tech} Software Engineer. Work on scalable solutions and collaborate with talented engineers.",
                "apply_link": "https://example.com/apply",
                "posted_date": datetime.now(timezone.utc).isoformat(),
                "is_remote": False,
                "matched_technology": tech,
                "source": "Sample"
            },
            {
                "job_id": f"sample_3_{tech.lower()}",
                "title": f"Lead {tech} Engineer",
                "company": "StartupX",
                "company_logo": "https://ui-avatars.com/api/?name=SX&background=ec4899&color=fff",
                "location": "Austin, TX",
                "employment_type": "Full-time",
                "salary_min": 160000,
                "salary_max": 220000,
                "description": f"Looking for a Lead {tech} Engineer to drive technical decisions and mentor junior developers. Remote-friendly position.",
                "apply_link": "https://example.com/apply",
                "posted_date": datetime.now(timezone.utc).isoformat(),
                "is_remote": True,
                "matched_technology": tech,
                "source": "Sample"
            },
            {
                "job_id": f"sample_4_{tech.lower()}",
                "title": f"{tech} Backend Developer",
                "company": "CloudScale Systems",
                "company_logo": "https://ui-avatars.com/api/?name=CS&background=14b8a6&color=fff",
                "location": "Seattle, WA",
                "employment_type": "Full-time",
                "salary_min": 140000,
                "salary_max": 190000,
                "description": f"Build robust backend services using {tech}. Experience with cloud platforms and microservices architecture preferred.",
                "apply_link": "https://example.com/apply",
                "posted_date": datetime.now(timezone.utc).isoformat(),
                "is_remote": True,
                "matched_technology": tech,
                "source": "Sample"
            },
            {
                "job_id": f"sample_5_{tech.lower()}",
                "title": f"Full Stack {tech} Developer",
                "company": "Digital Solutions Co",
                "company_logo": "https://ui-avatars.com/api/?name=DS&background=f59e0b&color=fff",
                "location": "Remote",
                "employment_type": "Full-time",
                "salary_min": 120000,
                "salary_max": 170000,
                "description": f"Full stack role focusing on {tech} for backend with React frontend. Great benefits and flexible work schedule.",
                "apply_link": "https://example.com/apply",
                "posted_date": datetime.now(timezone.utc).isoformat(),
                "is_remote": True,
                "matched_technology": tech,
                "source": "Sample"
            }
        ]
        return sample_jobs
    
    return {
        "recommendations": get_sample_jobs(primary_tech),
        "total": 5,
        "based_on": {
            "primary_technology": primary_tech,
            "sub_technologies": sub_techs
        },
        "data_source": "sample_fallback",
        "message": "Showing sample jobs. Live job scraping temporarily unavailable."
    }

@api_router.get("/live-jobs/{job_id}")
async def get_live_job_details(job_id: str, request: Request):
    """
    Get detailed information about a specific job.
    """
    await get_current_user(request)
    
    # Get API keys at request time
    rapidapi_key = os.environ.get('RAPIDAPI_KEY')
    rapidapi_host = os.environ.get('RAPIDAPI_HOST', 'jsearch.p.rapidapi.com')
    
    if not rapidapi_key:
        raise HTTPException(status_code=500, detail="JSearch API key not configured")
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as http_client:
            response = await http_client.get(
                "https://jsearch.p.rapidapi.com/job-details",
                params={"job_id": job_id},
                headers={
                    "X-RapidAPI-Key": rapidapi_key,
                    "X-RapidAPI-Host": rapidapi_host
                }
            )
            
            if response.status_code != 200:
                raise HTTPException(status_code=response.status_code, detail="Failed to fetch job details")
            
            data = response.json()
            jobs = data.get("data", [])
            
            if not jobs:
                raise HTTPException(status_code=404, detail="Job not found")
            
            job = jobs[0]
            
            return {
                "job_id": job.get("job_id"),
                "title": job.get("job_title"),
                "company": job.get("employer_name"),
                "company_logo": job.get("employer_logo"),
                "company_website": job.get("employer_website"),
                "location": job.get("job_city", "") + (", " + job.get("job_state", "") if job.get("job_state") else ""),
                "country": job.get("job_country"),
                "description": job.get("job_description"),
                "employment_type": job.get("job_employment_type"),
                "is_remote": job.get("job_is_remote", False),
                "apply_link": job.get("job_apply_link"),
                "posted_at": job.get("job_posted_at_datetime_utc"),
                "expires_at": job.get("job_offer_expiration_datetime_utc"),
                "salary_min": job.get("job_min_salary"),
                "salary_max": job.get("job_max_salary"),
                "salary_currency": job.get("job_salary_currency"),
                "salary_period": job.get("job_salary_period"),
                "source": job.get("job_publisher"),
                "highlights": job.get("job_highlights", {}),
                "required_skills": job.get("job_required_skills", []),
                "benefits": job.get("job_benefits", []),
                "qualifications": job.get("job_highlights", {}).get("Qualifications", []),
                "responsibilities": job.get("job_highlights", {}).get("Responsibilities", []),
            }
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting job details: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get job details: {str(e)}")

# ============ LIVE JOBS 2 (LinkedIn Job Search API) ============

# Configuration for Live Jobs 2 API
LIVEJOBS2_API_KEY = os.environ.get('LIVEJOBS2_API_KEY', '')
LIVEJOBS2_API_HOST = os.environ.get('LIVEJOBS2_API_HOST', 'linkedin-job-search-api.p.rapidapi.com')

@api_router.get("/live-jobs-2/search")
async def search_live_jobs_2(
    request: Request,
    query: Optional[str] = None,
    location: str = "United States",
    employment_type: Optional[str] = None,
    page: int = 1
):
    """
    Search jobs from LinkedIn Job Search API.
    """
    user = await get_current_user(request)
    
    # Get API keys at request time
    api_key = os.environ.get('LIVEJOBS2_API_KEY')
    api_host = os.environ.get('LIVEJOBS2_API_HOST', 'linkedin-job-search-api.p.rapidapi.com')
    
    if not api_key:
        return {
            "jobs": [],
            "total": 0,
            "page": page,
            "message": "Live Jobs 2 API not configured. Please provide LIVEJOBS2_API_KEY in backend/.env"
        }
    
    search_query = query or user.get('primary_technology', 'Software Developer')
    offset = (page - 1) * 10
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as http_client:
            # LinkedIn Job Search API endpoint
            response = await http_client.get(
                f"https://{api_host}/active-jb-24h",
                params={
                    "limit": 10,
                    "offset": offset,
                    "title_filter": f'"{search_query}"',
                    "location_filter": f'"{location}"',
                    "description_type": "text"
                },
                headers={
                    "X-RapidAPI-Key": api_key,
                    "X-RapidAPI-Host": api_host
                }
            )
            
            if response.status_code != 200:
                logger.warning(f"Live Jobs 2 API returned status {response.status_code}")
                return {"jobs": [], "total": 0, "page": page}
            
            jobs_data = response.json()
            if not isinstance(jobs_data, list):
                jobs_data = jobs_data.get("data", [])
            
            jobs = []
            for job in jobs_data:
                # Parse location from locations_derived
                location_str = ""
                if job.get("locations_derived"):
                    location_str = job["locations_derived"][0] if job["locations_derived"] else ""
                elif job.get("cities_derived"):
                    location_str = job["cities_derived"][0]
                    if job.get("regions_derived"):
                        location_str += f", {job['regions_derived'][0]}"
                
                # Parse salary
                salary_min = None
                salary_max = None
                salary_currency = "USD"
                salary_period = ""
                if job.get("salary_raw") and job["salary_raw"].get("value"):
                    salary_value = job["salary_raw"]["value"]
                    salary_min = salary_value.get("minValue")
                    salary_max = salary_value.get("maxValue")
                    salary_currency = job["salary_raw"].get("currency", "USD")
                    salary_period = salary_value.get("unitText", "")
                
                # Parse employment type
                emp_type = ""
                if job.get("employment_type"):
                    emp_type = job["employment_type"][0] if isinstance(job["employment_type"], list) else job["employment_type"]
                
                description = job.get("description_text", "")
                
                jobs.append({
                    "job_id": job.get("id", str(uuid.uuid4())),
                    "title": job.get("title", ""),
                    "company": job.get("organization", ""),
                    "company_logo": job.get("organization_logo", ""),
                    "location": location_str,
                    "country": job["countries_derived"][0] if job.get("countries_derived") else "",
                    "description": description[:500] + "..." if len(description) > 500 else description,
                    "full_description": description,
                    "employment_type": emp_type,
                    "is_remote": job.get("remote_derived", False),
                    "apply_link": job.get("url") or job.get("external_apply_url", ""),
                    "posted_at": job.get("date_posted", ""),
                    "salary_min": salary_min,
                    "salary_max": salary_max,
                    "salary_currency": salary_currency,
                    "salary_period": salary_period,
                    "source": "LinkedIn",
                    "required_skills": [],
                    "industry": job.get("linkedin_org_industry", ""),
                    "company_size": job.get("linkedin_org_size", ""),
                })
            
            return {
                "jobs": jobs,
                "total": len(jobs),
                "page": page
            }
            
    except Exception as e:
        logger.error(f"Error searching Live Jobs 2: {str(e)}")
        return {"jobs": [], "total": 0, "page": page, "error": str(e)}

@api_router.get("/live-jobs-2/recommendations")
async def get_live_jobs_2_recommendations(request: Request):
    """
    Get job recommendations from LinkedIn Job Search API based on user's profile.
    Requires user to have primary_technology set.
    """
    user = await get_current_user(request)
    
    # Check if user has required profile fields
    if not user.get('primary_technology'):
        return {
            "recommendations": [],
            "message": "Please update your profile with Primary Technology to get personalized job recommendations.",
            "requires_profile_update": True,
            "missing_fields": ["primary_technology"]
        }
    
    primary_tech = user.get('primary_technology')
    sub_techs = user.get('sub_technologies', [])
    
    # Sample jobs fallback when API is unavailable
    def get_sample_jobs_linkedin(tech):
        sample_jobs = [
            {
                "job_id": f"linkedin_sample_1_{tech.lower()}",
                "title": f"Senior {tech} Engineer",
                "company": "Meta Platforms",
                "company_logo": "https://ui-avatars.com/api/?name=MP&background=0866ff&color=fff",
                "location": "Menlo Park, CA",
                "employment_type": "Full-time",
                "salary_min": 180000,
                "salary_max": 250000,
                "description": f"Join Meta as a Senior {tech} Engineer. Work on products that connect billions of people worldwide.",
                "apply_link": "https://careers.meta.com",
                "posted_date": datetime.now(timezone.utc).isoformat(),
                "is_remote": True,
                "matched_technology": tech,
                "source": "LinkedIn (Sample)",
                "industry": "Technology"
            },
            {
                "job_id": f"linkedin_sample_2_{tech.lower()}",
                "title": f"{tech} Developer - Remote",
                "company": "Google LLC",
                "company_logo": "https://ui-avatars.com/api/?name=G&background=4285f4&color=fff",
                "location": "Mountain View, CA",
                "employment_type": "Full-time",
                "salary_min": 170000,
                "salary_max": 240000,
                "description": f"Google is hiring {tech} Developers to build next-generation products. Remote-friendly position with excellent benefits.",
                "apply_link": "https://careers.google.com",
                "posted_date": datetime.now(timezone.utc).isoformat(),
                "is_remote": True,
                "matched_technology": tech,
                "source": "LinkedIn (Sample)",
                "industry": "Technology"
            },
            {
                "job_id": f"linkedin_sample_3_{tech.lower()}",
                "title": f"Staff {tech} Engineer",
                "company": "Amazon",
                "company_logo": "https://ui-avatars.com/api/?name=A&background=ff9900&color=fff",
                "location": "Seattle, WA",
                "employment_type": "Full-time",
                "salary_min": 200000,
                "salary_max": 300000,
                "description": f"Amazon Web Services is looking for a Staff {tech} Engineer to drive technical initiatives and mentor teams.",
                "apply_link": "https://amazon.jobs",
                "posted_date": datetime.now(timezone.utc).isoformat(),
                "is_remote": False,
                "matched_technology": tech,
                "source": "LinkedIn (Sample)",
                "industry": "Cloud Computing"
            },
            {
                "job_id": f"linkedin_sample_4_{tech.lower()}",
                "title": f"{tech} Software Architect",
                "company": "Microsoft",
                "company_logo": "https://ui-avatars.com/api/?name=MS&background=00a4ef&color=fff",
                "location": "Redmond, WA",
                "employment_type": "Full-time",
                "salary_min": 190000,
                "salary_max": 280000,
                "description": f"Design and implement {tech} solutions at scale. Lead architectural decisions for cloud-native applications.",
                "apply_link": "https://careers.microsoft.com",
                "posted_date": datetime.now(timezone.utc).isoformat(),
                "is_remote": True,
                "matched_technology": tech,
                "source": "LinkedIn (Sample)",
                "industry": "Technology"
            },
            {
                "job_id": f"linkedin_sample_5_{tech.lower()}",
                "title": f"Principal {tech} Engineer",
                "company": "Netflix",
                "company_logo": "https://ui-avatars.com/api/?name=N&background=e50914&color=fff",
                "location": "Los Gatos, CA",
                "employment_type": "Full-time",
                "salary_min": 220000,
                "salary_max": 350000,
                "description": f"Netflix is seeking a Principal {tech} Engineer to lead streaming infrastructure development.",
                "apply_link": "https://jobs.netflix.com",
                "posted_date": datetime.now(timezone.utc).isoformat(),
                "is_remote": True,
                "matched_technology": tech,
                "source": "LinkedIn (Sample)",
                "industry": "Entertainment"
            }
        ]
        return sample_jobs
    
    api_key = os.environ.get('LIVEJOBS2_API_KEY')
    api_host = os.environ.get('LIVEJOBS2_API_HOST', 'linkedin-job-search-api.p.rapidapi.com')
    
    if not api_key:
        return {
            "recommendations": get_sample_jobs_linkedin(primary_tech),
            "user_technology": primary_tech,
            "api_status": "Sample data - API not configured"
        }
    
    # Get user's technologies
    user_technologies = [primary_tech]
    if sub_techs:
        user_technologies.extend(sub_techs[:2])
    
    all_recommendations = []
    api_error = None
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as http_client:
            for tech in user_technologies[:2]:  # Limit to 2 searches
                response = await http_client.get(
                    f"https://{api_host}/active-jb-24h",
                    params={
                        "limit": 5,
                        "offset": 0,
                        "title_filter": f'"{tech}"',
                        "location_filter": '"United States"',
                        "description_type": "text"
                    },
                    headers={
                        "X-RapidAPI-Key": api_key,
                        "X-RapidAPI-Host": api_host
                    }
                )
                
                if response.status_code == 200:
                    jobs_data = response.json()
                    
                    # Check for quota exceeded error
                    if isinstance(jobs_data, dict) and 'message' in jobs_data:
                        api_error = jobs_data.get('message', 'API error')
                        break
                    
                    if not isinstance(jobs_data, list):
                        jobs_data = jobs_data.get("data", [])
                    
                    for job in jobs_data[:5]:
                        # Parse location
                        location_str = ""
                        if job.get("locations_derived"):
                            location_str = job["locations_derived"][0] if job["locations_derived"] else ""
                        elif job.get("cities_derived"):
                            location_str = job["cities_derived"][0]
                            if job.get("regions_derived"):
                                location_str += f", {job['regions_derived'][0]}"
                        
                        # Parse salary
                        salary_min = None
                        salary_max = None
                        salary_currency = "USD"
                        salary_period = ""
                        if job.get("salary_raw") and job["salary_raw"].get("value"):
                            salary_value = job["salary_raw"]["value"]
                            salary_min = salary_value.get("minValue")
                            salary_max = salary_value.get("maxValue")
                            salary_currency = job["salary_raw"].get("currency", "USD")
                            salary_period = salary_value.get("unitText", "")
                        
                        # Parse employment type
                        emp_type = ""
                        if job.get("employment_type"):
                            emp_type = job["employment_type"][0] if isinstance(job["employment_type"], list) else job["employment_type"]
                        
                        description = job.get("description_text", "")
                        
                        all_recommendations.append({
                            "job_id": job.get("id", str(uuid.uuid4())),
                            "title": job.get("title", ""),
                            "company": job.get("organization", ""),
                            "company_logo": job.get("organization_logo", ""),
                            "location": location_str,
                            "country": job["countries_derived"][0] if job.get("countries_derived") else "",
                            "description": description[:500] + "..." if len(description) > 500 else description,
                            "full_description": description,
                            "employment_type": emp_type,
                            "is_remote": job.get("remote_derived", False),
                            "apply_link": job.get("url") or job.get("external_apply_url", ""),
                            "posted_at": job.get("date_posted", ""),
                            "salary_min": salary_min,
                            "salary_max": salary_max,
                            "salary_currency": salary_currency,
                            "salary_period": salary_period,
                            "source": "LinkedIn",
                            "required_skills": [],
                            "matched_technology": tech,
                            "industry": job.get("linkedin_org_industry", ""),
                            "company_size": job.get("linkedin_org_size", ""),
                        })
        
        # Check if we got API error or no results - fall back to sample data
        if api_error or not all_recommendations:
            return {
                "recommendations": get_sample_jobs_linkedin(primary_tech),
                "user_technology": primary_tech,
                "api_status": f"Using sample data - {api_error or 'No jobs found from API'}"
            }
        
        return {
            "recommendations": all_recommendations[:10],
            "user_technology": primary_tech
        }
        
    except Exception as e:
        logger.error(f"Error getting Live Jobs 2 recommendations: {str(e)}")
        return {
            "recommendations": get_sample_jobs_linkedin(primary_tech),
            "user_technology": primary_tech,
            "api_status": f"Using sample data - API error: {str(e)}"
        }

@api_router.get("/live-jobs-2/{job_id}")
async def get_live_job_2_details(job_id: str, request: Request):
    """
    Get detailed information about a specific job from LinkedIn Job Search API.
    """
    await get_current_user(request)
    
    # For LinkedIn Job Search API, job details are typically included in search results
    # This endpoint can be extended if the API provides a dedicated details endpoint
    return {
        "job_id": job_id,
        "message": "Job details are included in search results. Use the search or recommendations endpoint."
    }

# ============ AUTO-APPLY FEATURE ============

@api_router.get("/auto-apply/settings")
async def get_auto_apply_settings(request: Request):
    """Get user's auto-apply settings"""
    user = await get_current_user(request)
    
    settings = await db.auto_apply_settings.find_one(
        {"user_id": user["user_id"]},
        {"_id": 0}
    )
    
    if not settings:
        # Return default settings
        return {
            "user_id": user["user_id"],
            "enabled": False,
            "resume_id": "",
            "job_keywords": [user.get("primary_technology", "Software Developer")],
            "locations": ["United States"],
            "employment_types": ["FULL_TIME"],
            "min_salary": None,
            "max_applications_per_day": 10,
            "auto_tailor_resume": True,
            "last_run": None,
            "total_applications": 0
        }
    
    return settings

@api_router.post("/auto-apply/settings")
async def update_auto_apply_settings(data: AutoApplySettingsUpdate, request: Request):
    """Update user's auto-apply settings"""
    user = await get_current_user(request)
    
    # Get existing settings or create new
    existing = await db.auto_apply_settings.find_one({"user_id": user["user_id"]})
    
    update_data = {k: v for k, v in data.dict().items() if v is not None}
    update_data["user_id"] = user["user_id"]
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    if existing:
        await db.auto_apply_settings.update_one(
            {"user_id": user["user_id"]},
            {"$set": update_data}
        )
    else:
        update_data["created_at"] = datetime.now(timezone.utc).isoformat()
        update_data["last_run"] = None
        update_data["total_applications"] = 0
        await db.auto_apply_settings.insert_one(update_data)
    
    return {"message": "Settings updated successfully", "settings": update_data}

@api_router.post("/auto-apply/auto-fill-settings")
async def auto_fill_settings_from_profile(request: Request):
    """Auto-fill auto-apply settings from user's profile and resume"""
    user = await get_current_user(request)
    user_id = user["user_id"]
    
    # Get user's primary resume
    primary_resume = await db.resumes.find_one(
        {"user_id": user_id, "is_primary": True},
        {"_id": 0}
    )
    
    # If no primary, get the most recent resume
    if not primary_resume:
        cursor = db.resumes.find(
            {"user_id": user_id},
            {"_id": 0}
        ).sort("created_at", -1).limit(1)
        resumes_list = await cursor.to_list(1)
        primary_resume = resumes_list[0] if resumes_list else None
    
    # Build settings from profile
    job_keywords = []
    if user.get("primary_technology"):
        job_keywords.append(user["primary_technology"])
    if user.get("sub_technologies"):
        # Add top 2 sub-technologies
        job_keywords.extend(user["sub_technologies"][:2])
    
    # Default keywords if none from profile
    if not job_keywords:
        job_keywords = ["Software Developer"]
    
    # Get locations from profile
    locations = []
    if user.get("location_preferences"):
        locations = list(user["location_preferences"])
    elif user.get("location"):
        locations = [user["location"]]
    
    # Default to remote in US
    if not locations:
        locations = ["Remote, United States"]
    
    # Get job type preferences
    job_types = list(user.get("job_type_preferences", ["Remote"]) or ["Remote"])
    
    # Get salary preferences
    min_salary = user.get("salary_min")
    max_salary = user.get("salary_max")
    
    # Build auto-fill settings
    auto_filled_settings = {
        "user_id": user_id,
        "job_keywords": job_keywords[:5],  # Limit to 5 keywords
        "locations": locations[:3],  # Limit to 3 locations
        "job_types": job_types,
        "salary_min": min_salary,
        "salary_max": max_salary,
        "remote_only": "Remote" in job_types,
        "resume_id": primary_resume.get("resume_id") if primary_resume else None,
        "max_applications_per_day": 10,
        "auto_tailor_resume": True,
        "enabled": True,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    # Update or create settings
    existing = await db.auto_apply_settings.find_one({"user_id": user_id})
    if existing:
        await db.auto_apply_settings.update_one(
            {"user_id": user_id},
            {"$set": auto_filled_settings}
        )
    else:
        insert_data = auto_filled_settings.copy()
        insert_data["created_at"] = datetime.now(timezone.utc).isoformat()
        insert_data["last_run"] = None
        insert_data["total_applications"] = 0
        await db.auto_apply_settings.insert_one(insert_data)
    
    return {
        "message": "Settings auto-filled from profile successfully",
        "settings": auto_filled_settings
    }

@api_router.post("/auto-apply/toggle")
async def toggle_auto_apply(request: Request):
    """Toggle auto-apply on/off"""
    user = await get_current_user(request)
    
    settings = await db.auto_apply_settings.find_one({"user_id": user["user_id"]})
    
    if settings:
        new_status = not settings.get("enabled", False)
        await db.auto_apply_settings.update_one(
            {"user_id": user["user_id"]},
            {"$set": {"enabled": new_status, "updated_at": datetime.now(timezone.utc).isoformat()}}
        )
    else:
        new_status = True
        await db.auto_apply_settings.insert_one({
            "user_id": user["user_id"],
            "enabled": True,
            "resume_id": "",
            "job_keywords": [user.get("primary_technology", "Software Developer")],
            "locations": ["United States"],
            "employment_types": ["FULL_TIME"],
            "min_salary": None,
            "max_applications_per_day": 10,
            "auto_tailor_resume": True,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "last_run": None,
            "total_applications": 0
        })
    
    return {"enabled": new_status, "message": f"Auto-apply {'enabled' if new_status else 'disabled'}"}

@api_router.get("/auto-apply/history")
async def get_auto_apply_history(request: Request, limit: int = 50):
    """Get user's auto-apply history"""
    user = await get_current_user(request)
    
    history = await db.auto_applications.find(
        {"user_id": user["user_id"]},
        {"_id": 0}
    ).sort("applied_at", -1).limit(limit).to_list(limit)
    
    return {"history": history, "total": len(history)}

@api_router.post("/auto-apply/run")
async def run_auto_apply(request: Request):
    """
    Manually trigger auto-apply process.
    Fetches jobs using the system's job scraper and applies to matching jobs.
    """
    user = await get_current_user(request)
    user_id = user["user_id"]
    
    # Get user's auto-apply settings
    settings = await db.auto_apply_settings.find_one({"user_id": user_id})
    
    if not settings:
        raise HTTPException(status_code=400, detail="Please configure auto-apply settings first")
    
    if not settings.get("resume_id"):
        raise HTTPException(status_code=400, detail="Please select a resume for auto-apply")
    
    # Check if auto-apply is enabled
    if not settings.get("enabled", False):
        raise HTTPException(status_code=400, detail="Auto-apply is disabled. Please enable it first.")
    
    # Get the user's resume
    resume = await db.resumes.find_one(
        {"resume_id": settings["resume_id"], "user_id": user_id},
        {"_id": 0}
    )
    
    if not resume:
        raise HTTPException(status_code=404, detail="Selected resume not found")
    
    # Check daily limit
    today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
    today_applications = await db.auto_applications.count_documents({
        "user_id": user_id,
        "applied_at": {"$gte": today_start.isoformat()}
    })
    
    max_daily = settings.get("max_applications_per_day", 10)
    remaining = max_daily - today_applications
    
    if remaining <= 0:
        return {
            "message": f"Daily limit of {max_daily} applications reached",
            "applied_count": 0,
            "applications": []
        }
    
    # Get job search parameters from settings or profile
    job_keywords = settings.get("job_keywords", [])
    locations = settings.get("locations", ["Remote, United States"])
    
    # If no keywords, use user's primary technology
    if not job_keywords and user.get("primary_technology"):
        job_keywords = [user.get("primary_technology")]
    
    # Default to common tech jobs if still empty
    if not job_keywords:
        job_keywords = ["Software Developer"]
    
    # Use our system's job scraper to fetch jobs
    from utils.job_scraper import MultiSourceJobScraper
    
    all_jobs = []
    scraper = MultiSourceJobScraper()
    
    for keyword in job_keywords[:3]:  # Limit to 3 keywords
        for location in locations[:2]:  # Limit to 2 locations
            try:
                scraped_jobs = await scraper.scrape_all_sources(
                    keyword, 
                    location, 
                    limit_per_source=5
                )
                all_jobs.extend(scraped_jobs)
            except Exception as e:
                logger.error(f"Error scraping jobs for {keyword} in {location}: {str(e)}")
    
    # Remove duplicates based on job title and company
    seen = set()
    unique_jobs = []
    for job in all_jobs:
        key = (job.get('title', ''), job.get('company', ''))
        if key not in seen:
            seen.add(key)
            unique_jobs.append(job)
    
    # Filter for remote jobs only (based on user preferences)
    job_type_prefs = user.get("job_type_preferences", [])
    if "Remote" in job_type_prefs:
        unique_jobs = [j for j in unique_jobs if 
            'remote' in j.get('location', '').lower() or
            'remote' in j.get('title', '').lower()]
    
    # Limit to remaining daily quota
    jobs_to_apply = unique_jobs[:remaining]
    
    if not jobs_to_apply:
        return {
            "message": "No matching jobs found. Try adjusting your search keywords.",
            "applied_count": 0,
            "applications": []
        }
    
    # Process each job application
    applications = []
    resume_content = resume.get('master_resume') or resume.get('tailored_content') or resume.get('original_content', '')
    
    for job in jobs_to_apply:
        try:
            job_id = job.get("id", str(uuid.uuid4()))
            job_title = job.get("title", "")
            company = job.get("company", "")
            description = job.get("description", "")[:2000]  # Limit description length
            location_str = job.get("location", "")
            apply_link = job.get("apply_link", "")
            
            # Parse salary
            salary_info = job.get("salary", "")
            
            tailored_content = resume_content
            keywords_extracted = ""
            
            # Auto-tailor resume if enabled
            if settings.get("auto_tailor_resume", True) and description:
                try:
                    system_message = """You are an expert ATS resume optimizer. Transform resumes to match job requirements while maintaining authenticity. Focus on:
1. Matching keywords from the job description
2. Highlighting relevant experience
3. Using industry-standard formatting
4. Quantifying achievements where possible"""
                    
                    chat = LlmChat(
                        api_key=EMERGENT_LLM_KEY,
                        session_id=f"auto_tailor_{job_id}_{uuid.uuid4().hex[:8]}",
                        system_message=system_message
                    ).with_model("openai", "gpt-5.2")
                    
                    # Extract keywords
                    keywords_prompt = f"""Extract the top 15 most important keywords from this job posting for ATS optimization:

Job Title: {job_title}
Company: {company}
Description: {description[:1500]}

Return ONLY a comma-separated list of keywords."""
                    
                    keywords_message = UserMessage(text=keywords_prompt)
                    keywords_extracted = await chat.send_message(keywords_message)
                    
                    # Tailor resume
                    tailor_prompt = f"""Tailor this resume for the following job. Make it ATS-friendly and incorporate relevant keywords.

JOB DETAILS:
Title: {job_title}
Company: {company}
Key Requirements: {description[:1000]}

KEYWORDS TO INCORPORATE: {keywords_extracted}

ORIGINAL RESUME:
{original_content}

Return ONLY the tailored resume content."""
                    
                    tailor_message = UserMessage(text=tailor_prompt)
                    tailored_content = await chat.send_message(tailor_message)
                    
                except Exception as e:
                    logger.error(f"Error tailoring resume for job {job_id}: {str(e)}")
                    tailored_content = original_content
            
            # Create auto-application record
            application_record = {
                "application_id": f"auto_{uuid.uuid4().hex[:12]}",
                "user_id": user_id,
                "job_id": job_id,
                "job_title": job_title,
                "company": company,
                "location": location_str,
                "salary_info": salary_info,
                "job_description": description[:500],
                "apply_link": apply_link,
                "resume_id": settings["resume_id"],
                "tailored_content": tailored_content,
                "keywords_extracted": keywords_extracted,
                "status": "ready_to_apply",
                "applied_at": datetime.now(timezone.utc).isoformat(),
                "source": job.get("source", "system_scraper"),
                "auto_applied": True,
                "ats_optimized": True if tailored_content != resume_content else False
            }
            
            await db.auto_applications.insert_one(application_record)
            
            # Also record in main applications collection
            await db.applications.insert_one({
                "application_id": application_record["application_id"],
                "user_id": user_id,
                "job_portal_id": job.get("source", "auto_apply"),
                "job_title": job_title,
                "company_name": company,
                "job_description": description[:500],
                "resume_id": settings["resume_id"],
                "cover_letter": "",
                "status": "ready_to_apply",
                "created_at": datetime.now(timezone.utc).isoformat(),
                "auto_applied": True,
                "apply_link": apply_link
            })
            
            applications.append({
                "application_id": application_record["application_id"],
                "job_id": job_id,
                "job_title": job_title,
                "company": company,
                "location": location_str,
                "apply_link": apply_link,
                "status": "ready_to_apply"
            })
            
        except Exception as e:
            logger.error(f"Error processing job {job.get('id')}: {str(e)}")
            continue
    
    # Update settings with last run info
    await db.auto_apply_settings.update_one(
        {"user_id": user_id},
        {
            "$set": {"last_run": datetime.now(timezone.utc).isoformat()},
            "$inc": {"total_applications": len(applications)}
        }
    )
    
    return {
        "message": f"Successfully processed {len(applications)} job applications",
        "applied_count": len(applications),
        "applications": applications,
        "remaining_today": remaining - len(applications)
    }

@api_router.get("/auto-apply/status")
async def get_auto_apply_status(request: Request):
    """Get current auto-apply status including today's progress"""
    user = await get_current_user(request)
    user_id = user["user_id"]
    
    settings = await db.auto_apply_settings.find_one(
        {"user_id": user_id},
        {"_id": 0}
    )
    
    if not settings:
        return {
            "enabled": False,
            "configured": False,
            "today_applications": 0,
            "max_daily": 10,
            "remaining": 10,
            "last_run": None,
            "total_applications": 0
        }
    
    # Count today's applications
    today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
    today_applications = await db.auto_applications.count_documents({
        "user_id": user_id,
        "applied_at": {"$gte": today_start.isoformat()}
    })
    
    max_daily = settings.get("max_applications_per_day", 10)
    
    return {
        "enabled": settings.get("enabled", False),
        "configured": bool(settings.get("resume_id")),
        "today_applications": today_applications,
        "max_daily": max_daily,
        "remaining": max(0, max_daily - today_applications),
        "last_run": settings.get("last_run"),
        "total_applications": settings.get("total_applications", 0),
        "resume_id": settings.get("resume_id", ""),
        "job_keywords": settings.get("job_keywords", []),
        "locations": settings.get("locations", [])
    }


# ============ SCHEDULER MANAGEMENT ENDPOINTS ============

@api_router.get("/scheduler/status")
async def get_scheduler_status():
    """Get the current status of the auto-apply scheduler."""
    jobs = scheduler.get_jobs()
    
    job_info = []
    for job in jobs:
        job_info.append({
            "id": job.id,
            "name": job.name,
            "next_run_time": job.next_run_time.isoformat() if job.next_run_time else None,
            "trigger": str(job.trigger)
        })
    
    return {
        "scheduler_running": scheduler.running,
        "jobs": job_info,
        "timezone": "UTC"
    }


@api_router.post("/scheduler/trigger")
async def trigger_scheduled_auto_apply(request: Request):
    """Manually trigger the scheduled auto-apply for testing (admin only)."""
    user = await get_current_user(request)
    
    # For now, allow any authenticated user to trigger their own
    # In production, you might want to add admin check
    
    # Run for just this user
    settings = await db.auto_apply_settings.find_one(
        {"user_id": user["user_id"], "enabled": True},
        {"_id": 0}
    )
    
    if not settings:
        raise HTTPException(status_code=400, detail="Auto-apply is not enabled for your account")
    
    # Process in background
    asyncio.create_task(process_auto_apply_for_user(user["user_id"], settings))
    
    return {
        "message": "Scheduled auto-apply triggered successfully",
        "note": "Processing will continue in the background"
    }


@api_router.get("/scheduler/logs")
async def get_scheduler_logs(request: Request, limit: int = 20):
    """Get scheduler run logs for the current user."""
    user = await get_current_user(request)
    
    logs = await db.scheduler_logs.find(
        {"user_id": user["user_id"]},
        {"_id": 0}
    ).sort("run_at", -1).limit(limit).to_list(limit)
    
    return {"logs": logs}


@api_router.post("/auto-apply/schedule-settings")
async def update_schedule_settings(request: Request):
    """Update the user's preferred schedule time for auto-apply."""
    user = await get_current_user(request)
    data = await request.json()
    
    preferred_hour = data.get("preferred_hour", 6)  # Default 6 AM UTC
    
    if not 0 <= preferred_hour <= 23:
        raise HTTPException(status_code=400, detail="Hour must be between 0 and 23")
    
    await db.auto_apply_settings.update_one(
        {"user_id": user["user_id"]},
        {"$set": {"preferred_schedule_hour": preferred_hour}},
        upsert=True
    )
    
    return {
        "message": f"Schedule preference updated to {preferred_hour}:00 UTC",
        "preferred_hour": preferred_hour
    }


# ============ TECHNOLOGY OPTIONS ============

@api_router.get("/technologies")
async def get_technologies():
    return {
        "primary": ["Java", "Python", "PHP", "AI", "React"],
        "sub_technologies": {
            "Java": ["Spring Boot", "Hibernate", "Maven", "Gradle", "JUnit", "Microservices"],
            "Python": ["Django", "Flask", "FastAPI", "NumPy", "Pandas", "TensorFlow"],
            "PHP": ["Laravel", "Symfony", "CodeIgniter", "WordPress", "Drupal"],
            "AI": ["Machine Learning", "Deep Learning", "NLP", "Computer Vision", "PyTorch", "Keras"],
            "React": ["Next.js", "Redux", "TypeScript", "GraphQL", "Tailwind CSS", "Material UI"]
        }
    }

# Root endpoint
@api_router.get("/")
async def root():
    return {"message": "AI Resume Tailor API", "version": "1.0.0"}

# Health check endpoint for Kubernetes
@app.get("/health")
async def health_check():
    """Health check endpoint for Kubernetes liveness and readiness probes."""
    return {"status": "healthy", "service": "careerquest-api"}

# CORS - Add before including router
# Note: allow_origins must not be ["*"] when allow_credentials=True
cors_origins = os.environ.get('CORS_ORIGINS', '')
if cors_origins and cors_origins != '*':
    origins_list = [origin.strip() for origin in cors_origins.split(',') if origin.strip()]
else:
    # Default origins for development and production
    origins_list = [
        "http://localhost:3000",
        "https://job-tailor-7.preview.emergentagent.com",
        "https://job-tailor-7.emergentagent.com",
        "https://job-tailor-7.emergent.host",
        # Custom domain
        "https://hireignitor.com",
        "https://www.hireignitor.com",
    ]

# Use regex pattern to allow all emergentagent domains and custom domains
# This is more robust than listing individual origins
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=origins_list,
    allow_origin_regex=r"https://.*\.(emergentagent\.com|emergent\.host)$|https://(www\.)?hireignitor\.com$",
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
    allow_headers=["*"],
    expose_headers=["*"],
    max_age=600,  # Cache preflight response for 10 minutes
)

# Logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ============ SCHEDULER FUNCTIONS ============

async def scheduled_auto_apply_for_all_users():
    """
    Scheduled task that runs daily to auto-apply for all users with enabled auto-apply.
    This function processes each user who has auto-apply enabled.
    """
    logger.info("Starting scheduled auto-apply job for all users...")
    
    try:
        # Find all users with auto-apply enabled
        enabled_settings = await db.auto_apply_settings.find(
            {"enabled": True},
            {"_id": 0}
        ).to_list(1000)
        
        if not enabled_settings:
            logger.info("No users have auto-apply enabled. Skipping.")
            return
        
        logger.info(f"Found {len(enabled_settings)} users with auto-apply enabled")
        
        for settings in enabled_settings:
            user_id = settings.get("user_id")
            if not user_id:
                continue
                
            try:
                await process_auto_apply_for_user(user_id, settings)
            except Exception as e:
                logger.error(f"Error processing auto-apply for user {user_id}: {str(e)}")
                continue
                
        logger.info("Scheduled auto-apply job completed")
        
    except Exception as e:
        logger.error(f"Error in scheduled auto-apply job: {str(e)}")


async def process_auto_apply_for_user(user_id: str, settings: dict):
    """
    Process auto-apply for a single user.
    """
    logger.info(f"Processing auto-apply for user: {user_id}")
    
    if not settings.get("resume_id"):
        logger.warning(f"User {user_id} has no resume selected. Skipping.")
        return
    
    # Get the user's resume
    resume = await db.resumes.find_one(
        {"resume_id": settings["resume_id"], "user_id": user_id},
        {"_id": 0}
    )
    
    if not resume:
        logger.warning(f"Resume not found for user {user_id}. Skipping.")
        return
    
    # Check daily limit
    today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
    today_applications = await db.auto_applications.count_documents({
        "user_id": user_id,
        "applied_at": {"$gte": today_start.isoformat()}
    })
    
    max_daily = settings.get("max_applications_per_day", 10)
    remaining = max_daily - today_applications
    
    if remaining <= 0:
        logger.info(f"User {user_id} has reached daily limit of {max_daily}. Skipping.")
        return
    
    # Fetch jobs from LinkedIn API
    api_key = os.environ.get('LIVEJOBS2_API_KEY')
    api_host = os.environ.get('LIVEJOBS2_API_HOST', 'linkedin-job-search-api.p.rapidapi.com')
    
    if not api_key:
        logger.error("Job search API not configured")
        return
    
    job_keywords = settings.get("job_keywords", ["Software Developer"])
    locations = settings.get("locations", ["United States"])
    
    all_jobs = []
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as http_client:
            for keyword in job_keywords[:2]:
                for location in locations[:2]:
                    response = await http_client.get(
                        f"https://{api_host}/active-jb-24h",
                        params={
                            "limit": 10,
                            "offset": 0,
                            "title_filter": f'"{keyword}"',
                            "location_filter": f'"{location}"',
                            "description_type": "text"
                        },
                        headers={
                            "X-RapidAPI-Key": api_key,
                            "X-RapidAPI-Host": api_host
                        }
                    )
                    
                    if response.status_code == 200:
                        jobs_data = response.json()
                        if isinstance(jobs_data, list):
                            all_jobs.extend(jobs_data)
    except Exception as e:
        logger.error(f"Error fetching jobs for user {user_id}: {str(e)}")
        return
    
    # Filter out already applied jobs
    applied_job_ids = await db.auto_applications.distinct(
        "job_id",
        {"user_id": user_id}
    )
    
    new_jobs = [j for j in all_jobs if j.get("id") not in applied_job_ids][:remaining]
    
    if not new_jobs:
        logger.info(f"No new jobs found for user {user_id}")
        return
    
    applications_count = 0
    original_content = resume.get('original_content', '')
    
    # Process each job
    for job in new_jobs:
        try:
            job_id = job.get("id", str(uuid.uuid4()))
            job_title = job.get("title", "")
            company = job.get("organization", "")
            description = job.get("description_text", "")[:2000]
            
            location_str = ""
            if job.get("locations_derived"):
                location_str = job["locations_derived"][0]
            elif job.get("cities_derived"):
                location_str = job["cities_derived"][0]
            
            salary_info = ""
            if job.get("salary_raw") and job["salary_raw"].get("value"):
                salary_value = job["salary_raw"]["value"]
                if salary_value.get("minValue") and salary_value.get("maxValue"):
                    salary_info = f"${salary_value['minValue']:,} - ${salary_value['maxValue']:,}"
            
            tailored_content = original_content
            keywords_extracted = ""
            
            # Auto-tailor resume if enabled
            if settings.get("auto_tailor_resume", True) and description:
                try:
                    system_message = """You are an expert ATS resume optimizer. Transform resumes to match job requirements while maintaining authenticity."""
                    
                    chat = LlmChat(
                        api_key=EMERGENT_LLM_KEY,
                        session_id=f"scheduled_auto_{job_id}_{uuid.uuid4().hex[:8]}",
                        system_message=system_message
                    ).with_model("openai", "gpt-5.2")
                    
                    keywords_prompt = f"""Extract the top 15 most important keywords from this job posting:

Job Title: {job_title}
Company: {company}
Description: {description[:1500]}

Return ONLY a comma-separated list of keywords."""
                    
                    keywords_message = UserMessage(text=keywords_prompt)
                    keywords_extracted = await chat.send_message(keywords_message)
                    
                    tailor_prompt = f"""Tailor this resume for the following job. Make it ATS-friendly.

JOB: {job_title} at {company}
Key Requirements: {description[:1000]}

KEYWORDS: {keywords_extracted}

RESUME:
{original_content}

Return ONLY the tailored resume content."""
                    
                    tailor_message = UserMessage(text=tailor_prompt)
                    tailored_content = await chat.send_message(tailor_message)
                    
                except Exception as e:
                    logger.error(f"Error tailoring resume for job {job_id}: {str(e)}")
                    tailored_content = original_content
            
            # Create auto-application record
            application_record = {
                "application_id": f"scheduled_{uuid.uuid4().hex[:12]}",
                "user_id": user_id,
                "job_id": job_id,
                "job_title": job_title,
                "company": company,
                "location": location_str,
                "salary_info": salary_info,
                "job_description": description[:500],
                "apply_link": job.get("url") or job.get("external_apply_url", ""),
                "resume_id": settings["resume_id"],
                "tailored_content": tailored_content,
                "keywords_extracted": keywords_extracted,
                "status": "ready_to_apply",
                "applied_at": datetime.now(timezone.utc).isoformat(),
                "source": "LinkedIn_Scheduled",
                "auto_applied": True,
                "scheduled": True
            }
            
            await db.auto_applications.insert_one(application_record)
            
            await db.applications.insert_one({
                "application_id": application_record["application_id"],
                "user_id": user_id,
                "job_portal_id": "linkedin_scheduled",
                "job_title": job_title,
                "company_name": company,
                "job_description": description[:500],
                "resume_id": settings["resume_id"],
                "cover_letter": "",
                "status": "ready_to_apply",
                "created_at": datetime.now(timezone.utc).isoformat(),
                "auto_applied": True,
                "scheduled": True,
                "apply_link": application_record["apply_link"]
            })
            
            applications_count += 1
            
        except Exception as e:
            logger.error(f"Error processing job {job.get('id')} for user {user_id}: {str(e)}")
            continue
    
    # Update settings with last run info
    await db.auto_apply_settings.update_one(
        {"user_id": user_id},
        {
            "$set": {
                "last_scheduled_run": datetime.now(timezone.utc).isoformat(),
                "last_run": datetime.now(timezone.utc).isoformat()
            },
            "$inc": {"total_applications": applications_count}
        }
    )
    
    # Log the scheduled run
    await db.scheduler_logs.insert_one({
        "log_id": str(uuid.uuid4()),
        "user_id": user_id,
        "run_type": "scheduled_auto_apply",
        "applications_count": applications_count,
        "run_at": datetime.now(timezone.utc).isoformat(),
        "status": "completed"
    })
    
    logger.info(f"Completed auto-apply for user {user_id}: {applications_count} applications")


# ============ EMAIL CENTER ROUTES ============

@api_router.get("/email-center/accounts")
async def get_email_accounts(current_user: dict = Depends(get_current_user)):
    """Get all connected email accounts for the user"""
    accounts = await db.email_accounts.find(
        {"user_id": current_user["user_id"]},
        {"_id": 0, "password": 0, "refresh_token": 0, "access_token": 0}
    ).to_list(None)
    return accounts

@api_router.post("/email-center/connect/imap")
async def connect_imap_email(
    data: EmailAccountConnect,
    current_user: dict = Depends(get_current_user)
):
    """Connect an email account via IMAP/SMTP"""
    if data.provider != "imap":
        raise HTTPException(status_code=400, detail="Use this endpoint for IMAP connections only")
    
    if not all([data.imap_host, data.smtp_host, data.email_address, data.password]):
        raise HTTPException(status_code=400, detail="IMAP/SMTP connection requires host, email, and password")
    
    # Check if account already exists
    existing = await db.email_accounts.find_one({
        "user_id": current_user["user_id"],
        "email_address": data.email_address
    })
    
    if existing:
        raise HTTPException(status_code=400, detail="Email account already connected")
    
    # Test IMAP connection
    import imaplib
    try:
        if data.use_ssl:
            mail = imaplib.IMAP4_SSL(data.imap_host, data.imap_port)
        else:
            mail = imaplib.IMAP4(data.imap_host, data.imap_port)
        mail.login(data.email_address, data.password)
        mail.logout()
    except Exception as e:
        logger.error(f"IMAP connection failed: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Failed to connect: {str(e)}")
    
    # Check if this is the first account (make it primary)
    account_count = await db.email_accounts.count_documents({"user_id": current_user["user_id"]})
    
    account_id = f"email_{uuid.uuid4().hex[:12]}"
    account_doc = {
        "account_id": account_id,
        "user_id": current_user["user_id"],
        "provider": "imap",
        "email_address": data.email_address,
        "imap_host": data.imap_host,
        "imap_port": data.imap_port,
        "smtp_host": data.smtp_host,
        "smtp_port": data.smtp_port,
        "password": data.password,  # In production, encrypt this
        "use_ssl": data.use_ssl,
        "is_connected": True,
        "is_primary": account_count == 0,
        "connected_at": datetime.now(timezone.utc).isoformat(),
        "last_sync": None
    }
    
    await db.email_accounts.insert_one(account_doc)
    
    # Return without sensitive data
    del account_doc["password"]
    if "_id" in account_doc:
        del account_doc["_id"]
    
    return {"message": "Email account connected successfully", "account": account_doc}

@api_router.post("/email-center/connect/gmail/init")
async def init_gmail_oauth(current_user: dict = Depends(get_current_user)):
    """Initialize Gmail OAuth flow - returns auth URL"""
    # For Gmail, we'll use the user's Google account if they logged in with Google
    # Otherwise, they need to authorize Gmail access separately
    
    # Check if user has Google OAuth tokens from login
    user = await db.users.find_one({"user_id": current_user["user_id"]}, {"_id": 0})
    
    if user and user.get("google_access_token"):
        # User already has Google auth, use that
        return {
            "status": "already_connected",
            "message": "You can use your Google login email for sending applications",
            "email": user.get("email")
        }
    
    # For now, inform user to use IMAP for Gmail
    return {
        "status": "use_imap",
        "message": "Please use Gmail App Password with IMAP settings",
        "instructions": {
            "imap_host": "imap.gmail.com",
            "imap_port": 993,
            "smtp_host": "smtp.gmail.com", 
            "smtp_port": 587,
            "note": "Enable 2FA and create an App Password at https://myaccount.google.com/apppasswords"
        }
    }

@api_router.post("/email-center/connect/outlook/init")
async def init_outlook_oauth(current_user: dict = Depends(get_current_user)):
    """Initialize Outlook OAuth flow info"""
    return {
        "status": "use_imap",
        "message": "Please use Outlook App Password with IMAP settings",
        "instructions": {
            "imap_host": "outlook.office365.com",
            "imap_port": 993,
            "smtp_host": "smtp.office365.com",
            "smtp_port": 587,
            "note": "Enable 2FA and create an App Password in your Microsoft account security settings"
        }
    }

@api_router.delete("/email-center/accounts/{account_id}")
async def disconnect_email_account(
    account_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Disconnect an email account"""
    result = await db.email_accounts.delete_one({
        "account_id": account_id,
        "user_id": current_user["user_id"]
    })
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Account not found")
    
    return {"message": "Email account disconnected"}

@api_router.put("/email-center/accounts/{account_id}/primary")
async def set_primary_account(
    account_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Set an email account as primary"""
    # First, unset all primary flags for this user
    await db.email_accounts.update_many(
        {"user_id": current_user["user_id"]},
        {"$set": {"is_primary": False}}
    )
    
    # Set the specified account as primary
    result = await db.email_accounts.update_one(
        {"account_id": account_id, "user_id": current_user["user_id"]},
        {"$set": {"is_primary": True}}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Account not found")
    
    return {"message": "Primary account updated"}

@api_router.get("/email-center/inbox")
async def get_inbox_messages(
    account_id: Optional[str] = None,
    limit: int = 20,
    current_user: dict = Depends(get_current_user)
):
    """Get inbox messages from connected email account"""
    # Find the account to use
    query = {"user_id": current_user["user_id"]}
    if account_id:
        query["account_id"] = account_id
    else:
        query["is_primary"] = True
    
    account = await db.email_accounts.find_one(query)
    
    if not account:
        return {"messages": [], "message": "No email account connected"}
    
    if account["provider"] == "imap":
        try:
            import imaplib
            import email
            from email.header import decode_header
            
            if account.get("use_ssl", True):
                mail = imaplib.IMAP4_SSL(account["imap_host"], account.get("imap_port", 993))
            else:
                mail = imaplib.IMAP4(account["imap_host"], account.get("imap_port", 143))
            
            mail.login(account["email_address"], account["password"])
            mail.select("INBOX")
            
            # Search for recent emails
            _, message_numbers = mail.search(None, "ALL")
            message_ids = message_numbers[0].split()[-limit:]  # Get last N emails
            
            messages = []
            for num in reversed(message_ids):
                _, msg_data = mail.fetch(num, "(RFC822)")
                email_body = msg_data[0][1]
                msg = email.message_from_bytes(email_body)
                
                # Decode subject
                subject, encoding = decode_header(msg["Subject"])[0]
                if isinstance(subject, bytes):
                    subject = subject.decode(encoding or "utf-8")
                
                # Decode from
                from_header = msg.get("From", "")
                from_name = ""
                from_email = from_header
                if "<" in from_header:
                    parts = from_header.split("<")
                    from_name = parts[0].strip().strip('"')
                    from_email = parts[1].strip(">")
                
                # Get body preview
                body_preview = ""
                if msg.is_multipart():
                    for part in msg.walk():
                        if part.get_content_type() == "text/plain":
                            try:
                                body_preview = part.get_payload(decode=True).decode()[:200]
                            except:
                                pass
                            break
                else:
                    try:
                        body_preview = msg.get_payload(decode=True).decode()[:200]
                    except:
                        pass
                
                # Check if from recruiter (simple heuristic)
                is_recruiter = any(keyword in from_email.lower() or keyword in subject.lower() 
                    for keyword in ["recruit", "talent", "hr", "hiring", "career", "job", "opportunity"])
                
                messages.append({
                    "message_id": num.decode(),
                    "from_email": from_email,
                    "from_name": from_name,
                    "to_email": account["email_address"],
                    "subject": subject or "(No Subject)",
                    "body_preview": body_preview.strip()[:200] if body_preview else "",
                    "received_at": msg.get("Date", ""),
                    "is_read": True,  # IMAP doesn't easily provide this
                    "is_recruiter": is_recruiter
                })
            
            mail.logout()
            
            # Update last sync time
            await db.email_accounts.update_one(
                {"account_id": account["account_id"]},
                {"$set": {"last_sync": datetime.now(timezone.utc).isoformat()}}
            )
            
            return {"messages": messages, "account_email": account["email_address"]}
            
        except Exception as e:
            logger.error(f"Failed to fetch inbox: {str(e)}")
            return {"messages": [], "error": str(e)}
    
    return {"messages": [], "message": "Unsupported provider"}

@api_router.post("/email-center/send")
async def send_email(
    data: SendEmailRequest,
    current_user: dict = Depends(get_current_user)
):
    """Send an email using connected account"""
    # Find the account to use
    query = {"user_id": current_user["user_id"]}
    if data.account_id:
        query["account_id"] = data.account_id
    else:
        query["is_primary"] = True
    
    account = await db.email_accounts.find_one(query)
    
    if not account:
        raise HTTPException(status_code=400, detail="No email account connected")
    
    if account["provider"] == "imap":
        try:
            import smtplib
            from email.mime.text import MIMEText
            from email.mime.multipart import MIMEMultipart
            
            msg = MIMEMultipart()
            msg["From"] = account["email_address"]
            msg["To"] = ", ".join(data.to_addresses)
            msg["Subject"] = data.subject
            
            if data.body_type == "html":
                msg.attach(MIMEText(data.body, "html"))
            else:
                msg.attach(MIMEText(data.body, "plain"))
            
            # Connect to SMTP
            if account.get("smtp_port") == 465:
                server = smtplib.SMTP_SSL(account["smtp_host"], account["smtp_port"])
            else:
                server = smtplib.SMTP(account["smtp_host"], account.get("smtp_port", 587))
                server.starttls()
            
            server.login(account["email_address"], account["password"])
            server.send_message(msg)
            server.quit()
            
            # Log the sent email
            await db.email_center_history.insert_one({
                "history_id": f"sent_{uuid.uuid4().hex[:12]}",
                "user_id": current_user["user_id"],
                "account_id": account["account_id"],
                "type": "sent",
                "to_addresses": data.to_addresses,
                "subject": data.subject,
                "body_preview": data.body[:200],
                "sent_at": datetime.now(timezone.utc).isoformat()
            })
            
            return {"message": "Email sent successfully"}
            
        except Exception as e:
            logger.error(f"Failed to send email: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to send email: {str(e)}")
    
    raise HTTPException(status_code=400, detail="Unsupported email provider")

@api_router.post("/email-center/ai/compose-application")
async def ai_compose_application(
    data: AIComposeRequest,
    current_user: dict = Depends(get_current_user)
):
    """AI composes a job application email"""
    # Get user's resume
    resume = await db.resumes.find_one(
        {"resume_id": data.resume_id, "user_id": current_user["user_id"]},
        {"_id": 0}
    )
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    user = await db.users.find_one({"user_id": current_user["user_id"]}, {"_id": 0})
    
    prompt = f"""You are an expert career coach helping a job seeker write a compelling job application email.

CANDIDATE INFORMATION:
Name: {user.get('name', 'Candidate')}
Current Skills: {', '.join(user.get('sub_technologies', []))}
Primary Technology: {user.get('primary_technology', 'Software Development')}

RESUME SUMMARY:
{resume.get('original_content', '')[:2000]}

JOB DETAILS:
Position: {data.job_title}
Company: {data.company_name}
Description: {data.job_description[:1500]}

TONE: {data.tone}

Write a professional job application email that:
1. Has a compelling subject line
2. Opens with a strong introduction mentioning the specific role
3. Highlights 2-3 relevant skills/experiences from the resume that match the job
4. Shows genuine interest in the company
5. Ends with a clear call to action
6. Is concise (under 250 words)

Format your response as:
SUBJECT: [Subject line]
---
[Email body]"""

    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            model="gpt-5.2"
        )
        response = await asyncio.to_thread(
            chat.send_message,
            UserMessage(content=prompt)
        )
        
        # Parse response
        response_text = response.content
        subject = ""
        body = response_text
        
        if "SUBJECT:" in response_text and "---" in response_text:
            parts = response_text.split("---", 1)
            subject = parts[0].replace("SUBJECT:", "").strip()
            body = parts[1].strip()
        
        return {
            "subject": subject or f"Application for {data.job_title} at {data.company_name}",
            "body": body,
            "recipient": data.recipient_email,
            "tone": data.tone
        }
        
    except Exception as e:
        logger.error(f"AI compose failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate email")

@api_router.post("/email-center/ai/draft-reply")
async def ai_draft_reply(
    data: AIReplyRequest,
    current_user: dict = Depends(get_current_user)
):
    """AI drafts a reply to a recruiter email"""
    user = await db.users.find_one({"user_id": current_user["user_id"]}, {"_id": 0})
    
    prompt = f"""You are helping a job seeker draft a professional reply to a recruiter email.

CANDIDATE: {user.get('name', 'Candidate')}
SKILLS: {', '.join(user.get('sub_technologies', []))}

ORIGINAL EMAIL:
Subject: {data.original_subject}
From: {data.sender_email}
Content:
{data.original_email}

ADDITIONAL CONTEXT FROM CANDIDATE: {data.context or 'None provided'}

TONE: {data.tone}

Draft a professional reply that:
1. Thanks them for reaching out
2. Shows interest and enthusiasm
3. Addresses any questions they asked
4. Suggests next steps (call, meeting, etc.)
5. Is concise and professional

Write ONLY the email body (no subject line needed for replies)."""

    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            model="gpt-5.2"
        )
        response = await asyncio.to_thread(
            chat.send_message,
            UserMessage(content=prompt)
        )
        
        return {
            "reply_body": response.content,
            "original_subject": data.original_subject,
            "reply_to": data.sender_email,
            "tone": data.tone
        }
        
    except Exception as e:
        logger.error(f"AI reply draft failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate reply")

@api_router.get("/email-center/settings")
async def get_email_center_settings(current_user: dict = Depends(get_current_user)):
    """Get email center settings"""
    settings = await db.email_center_settings.find_one(
        {"user_id": current_user["user_id"]},
        {"_id": 0}
    )
    
    if not settings:
        settings = {
            "user_id": current_user["user_id"],
            "auto_reply_enabled": False,
            "auto_apply_compose": True,
            "reply_approval_required": True,
            "signature": ""
        }
    
    return settings

@api_router.post("/email-center/settings")
async def update_email_center_settings(
    data: EmailCenterSettings,
    current_user: dict = Depends(get_current_user)
):
    """Update email center settings"""
    await db.email_center_settings.update_one(
        {"user_id": current_user["user_id"]},
        {"$set": {
            **data.dict(),
            "user_id": current_user["user_id"],
            "updated_at": datetime.now(timezone.utc).isoformat()
        }},
        upsert=True
    )
    
    return {"message": "Settings updated"}

@api_router.get("/email-center/history")
async def get_email_history(
    limit: int = 50,
    current_user: dict = Depends(get_current_user)
):
    """Get email send/receive history"""
    history = await db.email_center_history.find(
        {"user_id": current_user["user_id"]},
        {"_id": 0}
    ).sort("sent_at", -1).limit(limit).to_list(None)
    
    return history


# Include router AFTER all routes are defined
app.include_router(api_router)


@app.on_event("startup")
async def startup_event():
    """Start the scheduler when the app starts."""
    logger.info("Starting application and scheduler...")
    
    # Schedule the auto-apply job to run daily at 6:00 AM UTC
    scheduler.add_job(
        scheduled_auto_apply_for_all_users,
        CronTrigger(hour=6, minute=0),  # Run at 6:00 AM UTC daily
        id="daily_auto_apply",
        name="Daily Auto-Apply Job",
        replace_existing=True
    )
    
    scheduler.start()
    logger.info("Scheduler started. Daily auto-apply job scheduled for 6:00 AM UTC")


@app.on_event("shutdown")
async def shutdown_db_client():
    """Shutdown the scheduler and close DB connection."""
    logger.info("Shutting down scheduler...")
    scheduler.shutdown(wait=False)
    client.close()
